from docx import Document
import pdfplumber as plb
import os
import csv
import util
from library.TSP.FileProcess_TSP_S3S4 import Record,Processor,Picture

class Test:

    # def __init__(self, record):
        # self.results, self.details = self.compareResult(record)

    def compareResult(self, record):

        # 稳定性
        GSI_STAB = record.dict["稳定性"]
        # 设计围岩级别
        GSI_DSCR = record.dict["设计围岩级别"]
        # 推测围岩级别
        GSI_ESRG = record.dict["推测围岩级别"]
        # 桩号区间
        GSI_INTE = record.dict["桩号区间"]
        # 岩性
        GSI_LITH = record.dict["岩性"]
        # 风化程度
        GSI_WEA = record.dict["风化程度"]
        # 密度
        GSI_DENST = record.dict["密度"]
        # 纵波速度
        GSI_PWL = record.dict["纵波速度"]
        # 横波速度
        GSI_SWL = record.dict["横波速度"]
        # 泊松比
        GSI_PR = record.dict["泊松比"]
        # 动态杨氏模量
        GSI_DYM = record.dict["动态杨氏模量"]
        # 预报结果描述
        GSI_RESULT = record.dict["预报结果描述"]
        # 完整性
        GSI_ITGT = record.dict["完整性"]
        # 地下水
        GSI_WATER = record.dict["地下水"]
        # 特殊地质情况
        GSI_SGS = record.dict["特殊地质情况"]

        with open('./standard/TSP/standard_TSP_S3S4.csv', 'r', encoding='utf-8') as csv_file:
            csv_read = csv.reader(csv_file)
            results = []
            details = []
            flag = False
            # print(csv_read)
            for row in csv_read:
                if GSI_INTE == row[0]:
                    results.append(not flag)
                    if GSI_LITH != "" and (GSI_LITH in row[1] or row[1] in GSI_LITH):
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_LITH) + " \n标准输出：" + str(row[1]))

                    if GSI_WEA == row[2]:
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_WEA) + " \n标准输出：" + str(row[2]))

                    if GSI_DENST == row[3]:
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_DENST) + " \n标准输出：" + str(row[3]))

                    if GSI_PWL == row[4]:
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_PWL) + " \n标准输出：" + str(row[4]))

                    if GSI_SWL == row[5]:
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_SWL) + " \n标准输出：" + str(row[5]))

                    if GSI_PR == row[6]:
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_PR) + " \n标准输出：" + str(row[6]))

                    if GSI_DYM == row[7]:
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_DYM) + "\n标准输出：" + str(row[7]))

                    if GSI_RESULT.replace(" ", "") == row[8].replace(" ", ""):
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_RESULT) + " \n标准输出：" + str(row[8]))

                    if GSI_WATER == row[9]:
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_WATER) + " \n标准输出：" + str(row[9]))

                    if GSI_SGS.replace(" ", "") == row[10].replace(" ", ""):
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_SGS) + " \n标准输出：" + str(row[10]))

                    if GSI_STAB != "" and (GSI_STAB in row[11] or row[11] in GSI_STAB):
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_STAB) + " \n标准输出：" + str(row[11]))

                    if GSI_DSCR != "" and (
                            GSI_DSCR in row[12].replace(" ", "") or row[12].replace(" ", "") in GSI_DSCR):
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_DSCR) + " \n标准输出：" + str(row[12]))

                    if GSI_ESRG != "" and (GSI_ESRG in row[13] or row[13] in GSI_ESRG):
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_ESRG) + " \n标准输出：" + str(row[13]))

                    if GSI_ITGT.replace(" ", "") == row[14].replace(" ", ""):
                        results.append(not flag)
                        details.append("无")
                    else:
                        results.append(flag)
                        details.append("程序输出：" + str(GSI_ITGT) + " \n标准输出：" + str(row[14]))
                    break
        # print(results)
        # print(details)
        return results, details

class Execute:

    def run(self, input_path):
        process=Processor()
        files_to_process = set()
        files_to_transform = set()
        pdf_to_process = set()
        result_set = []
        count = 0
        for file in os.listdir(input_path):
            absolute_file_path = os.path.join(input_path, file)
            if file.endswith(".doc"):
                files_to_transform.add(absolute_file_path)
            elif file.endswith(".docx"):
                files_to_process.add(absolute_file_path)
            elif file.endswith(".pdf"):
                pdf_to_process.add(absolute_file_path)
        files_to_delete = util.batch_doc_to_docx(files_to_transform)
        files_to_process = files_to_process.union(files_to_delete)

        for file in files_to_process:
            docx = Document(file)
            records = list()
            table2 = docx.tables[-2]
            table3 = docx.tables[-1]
            conclusion = process.get_conclusion(docx)

            process.get_record_table3(records, table3)
            process.get_record_table2(records, table2)
            process.get_record_conclusion(records, conclusion)

            for record in records:
                test = Test()
                result, detail = test.compareResult(record)
                fault = self.filter(result, detail)
                if len(fault) != 0:
                    result_set.append(fault)
            count = count + 1
            # print(count)
            print("测试完成TSP_S3S4 Word文件" + file)

        for file in pdf_to_process:
            docx = Document()
            with plb.open(file) as pdf:
                tables = []
                content = ""
                for i in range(len(pdf.pages)):
                    table = pdf.pages[i].extract_tables()
                    content += pdf.pages[i].extract_text()

                    if not len(table) == 0:
                        tables.append(table)
                tb2 = tables[-2]
                table2 = docx.add_table(len(tb2[0]), len(tb2[0][0]))
                table2 = process.traverse_table(tb2, table2)
                tb3 = tables[-1]
                table3 = docx.add_table(len(tb3[0]), len(tb3[0][0]))
                table3 = process.traverse_table(tb3, table3)
                records = list()
                conclusion = process.get_pdf_conclusion(content)

                process.get_record_table3(records, table3)
                process.get_record_table2(records, table2)
                process.get_record_conclusion(records, conclusion)
            for record in records:
                test = Test()
                result, detail = test.compareResult(record)
                fault = self.filter(result, detail)
                result_set.append(fault)
            count = count + 1
            print("测试完成TSP_S3S4 PDF文件" + file)


        for file in files_to_delete:
            if os.path.exists(file):
                os.remove(file)
        # print("end count:")
        # print(count)
        return count, result_set

    def filter(self, results, details):
        if len(results) - 1 != len(details):
            print("长度不一致！")
            return
        fault = []
        for i in range(len(details)):
            if results[i + 1] == True:
                continue
            fault.append(details[i])
        return fault

if __name__ == "__main__":
    inputpath = "e:/study/is3/tsp2"
    a = Execute()
    count, result_set=a.run(inputpath)
    print(count)
    print(result_set)