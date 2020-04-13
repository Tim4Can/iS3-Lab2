# coding=utf-8
from docx import Document
import os
import csv
import re
import sys
sys.path.append("..")
from FileProcessBasic import FileProcessBasic
# from win32com.client import Dispatch

def check_output_file(output_path, header):
    if not os.path.exists(output_path):
        with open(output_path, "w", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, header)
            w.writeheader()

class Record:
    def __init__(self, docx):
        # self.face_station = None  # 掌子面桩号
        # self.number_interval = None  # 桩号区间
        # self.GPR_description = None  # 地质雷达描述
        # self.ground_water_description = None  # 地下水状态描述
        # self.ground_water_level = None  # 地下水对应等级
        # self.lithology = None  # 岩性
        # self.weathering_degree = None  # 风化程度
        # self.structure = None  # 结构构造
        # self.fault = None  # 断层
        # self.stability = None  # 稳定性
        # self.designed_rock_level = None  # 设计围岩级别
        # self.predict_rock_level = None  # 预报围岩级别
        self.dict = {
            "掌子面桩号": None,
            "桩号区间": None,
            "地质雷达描述": None,
            "地下水状态描述": None,
            "地下水对应等级": None,
            "岩性": None,
            "风化程度": None,
            "结构构造": None,
            "断层": None,
            "稳定性": None,
            "设计围岩级别": None,
            "预报围岩级别": None
        }
        self.get_face_station(docx.tables[2])
        self.para = self.locateParagraph(docx)
        self.findKeywords(self.para)


        
    def locateParagraph(self, docx):
        flag = 0
        para = ''
        for p in docx.paragraphs:
            if "6.2" in p.text and flag == 0:
                flag = 1
                continue
            if flag == 0:
                continue
            if flag == 1:
                para = para + p.text
        return para

    def findKeywords(self, para):
        # 掌子面桩号
        '''
        GSI_CHAI = ''
        for i in range (len(para)):
            if (para[i] == '+'):
                j = i
                while (para[j] != '（'):
                    j = j - 1
                while (para[j+1] != '～'):
                    GSI_CHAI = GSI_CHAI + (para[j+1])
                    j = j + 1
                break
        self.dict["掌子面桩号"] = GSI_CHAI
        '''
        # 桩号区间
        GSI_INTE = ''
        for i in range (len(para)):
            if (para[i] == '+'):
                j = i
                while (para[j] != '（'):
                    j = j - 1 
                while (para[j+1] != '）'):
                    GSI_INTE = GSI_INTE + (para[j+1])
                    j = j + 1
                break
        self.dict["桩号区间"] = GSI_INTE
        # 地质雷达描述
        GSI_GPR = ''
        for i in range (len(para)-2):
            if (para[i:i+3] == "电磁波"):
                j = i+3
                k = i+4
                times = 3
                while (para[j] != '，'):
                    GSI_GPR = para[j] + GSI_GPR
                    j = j-1
                while (para[k] != '，' or times != 0):
                    GSI_GPR = GSI_GPR + para[k]
                    k = k+1
                    if (para[k] == '，'):
                        times = times-1
                break
        self.dict["地质雷达描述"] = GSI_GPR

        # 地下水状态描述（此数据只体现在表格中，在段落中没有体现）
        GSI_WATE = ''
        print("地下水状态描述：此数据只体现在表格中，在段落中没有体现")

        # 地下水对应等级（此数据只体现在表格中，在段落中没有体现）
        GSI_WATG = ''
        print("地下水对应等级：此数据只体现在表格中，在段落中没有体现")
        
        # 岩性

        GSI_LITH = ''
        for i in range (len(para)-2):
            if (para[i:i+3] == "岩性为"):
                j = i+3
                while(para[j] != '，'):
                    GSI_LITH = GSI_LITH + para[j]
                    j = j+1
                break
        self.dict["岩性"] = GSI_LITH


        # 风化程度
        GSI_WEA = ''
        for i in range (len(para)-1):
            if (para[i:i+2] == "风化"):
                GSI_WEA = para[i-1:i+2]
                if(para[i-1] == "等"):
                    GSI_WEA = para[i-2:i+2]
                if(para[i-2] == "～"):
                    if(para[i-3] != "等"):
                        GSI_WEA = para[i-3:i+2]
                    else:
                        GSI_WEA = para[i-4:i+2]
                if(para[i-3] == "～"):
                    GSI_WEA = para[i-4:i+2]
        self.dict["风化程度"] = GSI_WEA

        # 结构构造
        GSI_STRU = ''
        for i in range (len(para)-1):
            if(para[i:i+2] == "结构"):
                j = i+1
                while (para[j] != '，'):
                    GSI_STRU = para[j] + GSI_STRU
                    j = j-1
                break
        self.dict["结构构造"] = GSI_STRU

        # 断层
        '''
        GSI_FAUL = '无'
        for i in range (len(para)-1):
            if (para[i:i+2] == "断层"):
                GSI_FAUL = "断层带"
        self.dict["断层"] = GSI_FAUL
        '''

        # 稳定性
        GSI_STAB = ''
        for i in range(len(para)-2):
            if (para[i:i+3] == "稳定性"):
                j = i+3
                t = 5
                while (para[j] != '，' or t != 0):
                    GSI_STAB = para[j]+GSI_STAB
                    j = j-1
                    if (para[j] == '，'):
                        t = t - 1
                break
        self.dict["稳定性"] = GSI_STAB

        # 设计围岩级别
        GSI_DSCR = ''
        for i in range(len(para)-6):
            if (para[i:i+7] == "设计围岩等级为"):
                GSI_DSCR = para[i+7]
                break
        self.dict["设计围岩级别"] = GSI_DSCR

        # 预报围岩级别
        GSI_PSRL = ''
        for i in range(len(para)-4):
            if (para[i:i+5] == "预判围岩为"):
                GSI_PSRL = para[i+5]
                break
        self.dict["预报围岩级别"] = GSI_PSRL
        

    def get_face_station(self, table):

        # 直接提取
        for i in range(len(table.rows)):
            tmp=list(table.rows[i].cells)
            cols=sorted(set(tmp),key=tmp.index)
            for j in range(len(cols)):
                if cols[j].text=='掌子面桩号' and j<len(cols)-1:
                    self.dict['掌子面桩号']=re.sub(u"\\（.*?）", "", cols[j+1].text)
                    break

        # 选择结果
        for i in range(len(table.rows)):
            text = table.cell(i,0).text 
            if text=='地下水状态':
                print("地下水状态")
                tmp=list(table.rows[i].cells)
                cols=sorted(set(tmp),key=tmp.index)
                for col in cols:
                    if col.text.find('√')>0:
                        print("yes1")
                        col.text=col.text.replace('√','')
                        self.dict['地下水对应等级']=col.text
                        break
                
            if text=='岩体出露状态':
                print("岩体出露状态")
                tmp=list(table.rows[i].cells)
                cols=sorted(set(tmp),key=tmp.index)
                for col in cols:
                    if col.text.find('√')>0:
                        print("yes2")
                        col.text=col.text.replace('√','')
                        self.dict['断层']=col.text
                        break 
                if self.dict['断层']==None:
                	self.dict['断层']='无'


        # return ""

    def to_string(self):
        res = ""
        for name, value in vars(self).items():
            if value is None:
                value = ""
            res = res + value + ","
        return res[: -1]

class Processor(FileProcessBasic):
    def save(self, output, record):
        output_path = os.path.join(output, "GPR_S1S2.csv")
        header = record.dict.keys()
        check_output_file(output_path, header)

        with open(output_path, "a+", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, record.dict.keys())
            w.writerow(record.dict)

    def run(self, inputpath, outputpath):
        transformed = False
        if inputpath.endswith("doc"):
            transformed = True
            try:
                # word = Dispatch('Word.Application')
                doc = Document(inputpath)
                # doc = documents.Open(inputpath)
                doc.SaveAs("{}x".format(inputpath), 12)
                doc.Close()
                # word.Quit()
            except IOError:
                print("读取文件异常：" + inputpath)
            inputpath = inputpath + "x"
        elif not inputpath.endswith("docx"):
            return

        docx = Document(inputpath)
        record = Record(docx)
        self.save(outputpath, record)

        print("提取完成" + inputpath)
        if transformed and os.path.exists(inputpath):
            os.remove(inputpath)


if __name__ == "__main__":
    test = Processor()
    inputpath = "D:/Death in TJU/Junior_2nd/iS3 Lab2/tasks/task3/GPRS1S2.docx"
    outputpath = "D:/Death in TJU/Junior_2nd/iS3 Lab2/tasks/task3"
    test.run(inputpath, outputpath)
