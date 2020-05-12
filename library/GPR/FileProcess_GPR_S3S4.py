from docx import Document
import os
import csv
import re
import xml.etree.cElementTree as ET
import pdfplumber as plb 
import fitz
from library.FileProcessBasic import FileProcessBasic
import util

def check_output_file(output_path, header):
    if not os.path.exists(output_path):
        with open(output_path, "w", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, header)
            w.writeheader()

class Record:
    def __init__(self, docx):
        name, GSI_INTE = self.get_cover(docx)
        GSI_CHAI, GSI_INTE = util.parse_GSI_CHAI_and_GSI_INTE(name, GSI_INTE)

        para_result, para_prediction = self.locate_paragraph(docx)
        GSI_GPR = self.get_GSI_GPR(para_result)

        GSI_LITH = self.get_GSI_LITH(para_prediction)
        GSI_FAUL = self.get_GSI_FAUL(para_prediction)
        GSI_STRU = self.get_GSI_STRU(para_prediction)
        GSI_STAB = self.get_GSI_STAB(para_prediction)
        GSI_WEA = self.get_GSI_WEA(para_prediction)

        GSI_DSCR = self.get_GSI_DSCR(docx.tables[3])
        GSI_PSRL = self.get_GSI_PSRL(docx.tables[3])

        GSI_WATE = self.get_GSI_WATE(docx.tables[4])
        GSI_WATG = self.get_GSI_WATG(docx.tables[4])

        self.dict = {
            "掌子面桩号": GSI_CHAI,
            "桩号区间": GSI_INTE,
            "地质雷达描述": GSI_GPR,
            "地下水状态描述": GSI_WATE,
            "地下水对应等级": GSI_WATG,
            "岩性": GSI_LITH,
            "风化程度": GSI_WEA,
            "结构构造": GSI_STRU,
            "断层": GSI_FAUL,
            "稳定性": GSI_STAB,
            "设计围岩级别": GSI_DSCR,
            "预报围岩级别": GSI_PSRL
        }

    def get_cover(self, docx):
        name, GSI_INTE = None, None
        for paragraph in docx.paragraphs:
            if paragraph.text.startswith("隧道名称："):
                name = paragraph.text.split("：")[1].strip()
            if paragraph.text.startswith("预报里程："):
                GSI_INTE = paragraph.text.split("：")[1].strip()
            if name is not None and GSI_INTE is not None:
                return name, GSI_INTE

    def locate_paragraph(self, docx):
        para_result = ""        # 6.2 探测结果
        para_prediction = ""    # 6.3 前方地质情况预测
        for i, p in enumerate(docx.paragraphs):
            if p.text.startswith("6.2"):
                i += 1
                p = docx.paragraphs[i]
                while not p.text.startswith("6.3"):
                    if not p.text.startswith("图"):
                        para_result += p.text
                    i += 1
                    p = docx.paragraphs[i]
                i -= 1
            elif p.text.startswith("6.3"):
                i += 1
                p = docx.paragraphs[i]
                while not p.text.startswith("7"):
                    para_prediction += p.text
                    i += 1
                    p = docx.paragraphs[i]
        return para_result, para_prediction

    # 掌子面桩号
    def get_GSI_CHAI_and_GSI_INTE(self, para):
        GSI_INTE = ''
        for i in range(len(para)):
            if (para[i] == '+'):
                j = i
                while (para[j + 1] != 'K'):
                    j = j - 1
                GSI_INTE = para[j + 1:j + 14]
                break
        GSI_CHAI = GSI_INTE.split("～")[0]
        return GSI_CHAI, GSI_INTE

    # 地质雷达描述
    def get_GSI_GPR(self, para):
        start = para.find("基本规律：") + 5
        end = para.find("。", start)
        GSI_GPR = para[start: end]
        if GSI_GPR is None:
            GSI_GPR = "无"
        return GSI_GPR

    # 岩性
    def get_GSI_LITH(self, para):
        keyword = "岩性："
        start = para.find(keyword) + len(keyword)
        end = para.find("。", start)
        GSI_LITH = para[start: end]

        keyword = "风化"
        GSI_LITH = GSI_LITH[GSI_LITH.find(keyword) + len(keyword):]
        if GSI_LITH is None:
            GSI_LITH = "无"
        return GSI_LITH

    # 风化程度
    def get_GSI_WEA(self, para):
        keyword = "岩性："
        start = para.find(keyword) + len(keyword)
        end = para.find("。", start)
        GSI_WEA = para[start: end]

        keyword = "风化"
        GSI_WEA = GSI_WEA[:GSI_WEA.find(keyword) + len(keyword)]
        if GSI_WEA is None:
            GSI_WEA = "无"
        return GSI_WEA

    # 结构构造
    def get_GSI_STRU(self, para):
        keyword = "结构构造："
        start = para.find(keyword) + len(keyword)
        end = para.find("。", start)
        GSI_STRU = para[start: end]
        if GSI_STRU == "":
            GSI_STRU = "无"
        return GSI_STRU

    # 稳定性
    def get_GSI_STAB(self, para):
        keyword = "稳定性分析："
        start = para.find(keyword) + len(keyword)
        end = para.find("。", start)
        GSI_STAB = para[start: end]

        if GSI_STAB is not None:
            GSI_STAB = GSI_STAB.split("，")[-1]
            key = "围岩"
            GSI_STAB = GSI_STAB[GSI_STAB.find(key) + len(key):]
        else:
            GSI_STAB = "无"
        return GSI_STAB

    # 设计围岩级别
    def get_GSI_DSCR(self, table):
        GSI_DSCR = table.rows[1].cells[3].text
        return GSI_DSCR

    # 预报围岩级别
    def get_GSI_PSRL(self, table):
        GSI_PSRL = table.rows[1].cells[-1].text
        return GSI_PSRL

    # 地下水状态描述
    def get_GSI_WATE(self, table):
        GSI_WATE = ""
        for row in table.rows:
            if row.cells[0].text.strip() == '地下水':
                watgs = set()
                for i in range(1, len(row.cells)):
                    cell = row.cells[i]
                    if "√" in cell.text:
                        watgs.add(cell.text.replace("√", "").strip())
                GSI_WATE = "~".join(watgs)
                if len(watgs) > 0:
                    break
        if GSI_WATE == "":
            GSI_WATE = "无"
        return GSI_WATE

    # 地下水对应等级
    def get_GSI_WATG(self, table):
        return "无"

    # 断层
    def get_GSI_FAUL(self, para):
        keyword = "断层破碎带："
        start = para.find(keyword) + len(keyword)
        end = para.find("。", start)
        GSI_FAUL = para[start: end]
        if GSI_FAUL == "":
            GSI_FAUL = "无"
        return GSI_FAUL

class Picture:
    def __init__(self, type_name, file_name, docx):
        self.file = file_name
        self.directory = self.parse_file(type_name, file_name)
        self.picture_ids = self.extract_graphs(docx)

    def extract_graphs(self, docx):
        ids = []
        flag = False
        for i, p in enumerate(docx.paragraphs):
            if not flag and p.text.replace(" ", "").strip() == "目录":
                flag = True
            if flag:
                root = ET.fromstring(p._p.xml)
                pic_str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
                pics = root.findall(pic_str)
                image_str = "*/{urn:schemas-microsoft-com:vml}shape/{urn:schemas-microsoft-com:vml}imagedata"
                for pic in pics:
                    pict = pic.findall(image_str)
                    if len(pict) > 0:
                        text = docx.paragraphs[i + 1].text
                        if not text.endswith("示意图"):
                            ids.append(pict[0].attrib[
                                           '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'])
        return ids

    def parse_file(self, type_name, file_name):
        stage = None
        match = re.search("\d{3}", file_name)
        if match is not None:
            span = match.span()
            stage = file_name[span[0]: span[1]]
            stage = str(int(stage))

        GSI_INTE = None
        match = re.search("K\d\+\d{3}[-～](K\d\+)?\d{3}", file_name)
        if match is not None:
            span = match.span()
            GSI_INTE = file_name[span[0]: span[1]]
            if "-" in GSI_INTE:
                GSI_INTE = GSI_INTE.split("-")
                pre = GSI_INTE[0][: 3]
                GSI_INTE[1] = pre + GSI_INTE[1]
                GSI_INTE = "~".join(GSI_INTE)

        prefix = util.map_prefix(util.parse_prefix(file_name))

        return type_name + prefix + stage + "期" + GSI_INTE


class RecordPDF:
	def __init__(self, file):
		# cover
		name, GSI_INTE = self.get_cover(file)
		GSI_CHAI, GSI_INTE = util.parse_GSI_CHAI_and_GSI_INTE(name, GSI_INTE)

		# paragraph
		para_result, para_prediction = self.locate_paragraph(file)
		GSI_GPR = self.get_GSI_GPR(para_result)
		GSI_STAB = self.get_GSI_STAB(para_prediction)
		GSI_STRU = self.get_GSI_STRU(para_prediction)
		GSI_LITH = self.get_GSI_LITH(para_prediction)
		GSI_WEA = self.get_GSI_WEA(para_prediction)

		# appendix table
		appendix = self.get_appendix(file)
		GSI_WATG = self.get_GSI_WATG(appendix)
		GSI_DSCR = self.get_GSI_DSCR(appendix)
		GSI_PSRL = self.get_GSI_PSRL(appendix)

		# 未实现
		GSI_WATE = "无"
		GSI_FAUL = "无"


		self.dict = {
			"掌子面桩号": GSI_CHAI,
			"桩号区间": GSI_INTE,
			"地质雷达描述": GSI_GPR,
			"地下水状态描述": GSI_WATE,
			"地下水对应等级": GSI_WATG,
			"岩性": GSI_LITH,
			"风化程度": GSI_WEA,
			"结构构造": GSI_STRU,
			"断层": GSI_FAUL,
			"稳定性": GSI_STAB,
			"设计围岩级别": GSI_DSCR,
			"预报围岩级别": GSI_PSRL
		}

		for key, value in self.dict.items():
			print(key + " " + value)


	# 获取封面
	def get_cover(self, file):
		name, GSI_INTE = None, None

		with plb.open(file) as pdf:
			text = pdf.pages[0].extract_text()
			if text:
				lines = text.splitlines()
				for line in lines:
					if line.startswith("隧道名称") or line.startswith("项目名称"):
						name = line.split("：")[1].strip()
					if line.startswith("预报里程"):
						GSI_INTE = line.split("：")[1].strip()
					if name is not None and GSI_INTE is not None:
						break

		return name, GSI_INTE

	# 获取段落
	def locate_paragraph(self, file):
		para_result, para_prediction = "", ""
		collect_result, collect_prediction = False, False

		with plb.open(file) as pdf:
			for page in pdf.pages:
				text = page.extract_text()
				lines = text.splitlines()
				for line in lines:
					line = line.strip()
					# 无意义，略过
					if line == "" or (line.startswith("第") and line.endswith("页")):
						continue

					# 提取探测结果
					if line.startswith("6.2"):
						collect_result = True
						continue
					elif line.startswith("6.3"):
						collect_result = False

					if collect_result == True:
						para_result += line
						continue

					# 提取前方地质情况预测
					if line.startswith("6.3"):
						collect_prediction = True
						continue
					elif line.startswith("7"):
						collect_prediction = False
						# 彻底结束
						break 

					if collect_prediction == True:
						para_prediction += line
						continue

		return para_result, para_prediction

	# 地质雷达描述
	def get_GSI_GPR(self, para):
		GSI_GPR = "无"

		try:
			start = para.find("基本规律：") + 5
			end = para.find("。", start)
			GSI_GPR = para[start: end]
			if GSI_GPR is None:
				GSI_GPR = "无"
			return GSI_GPR
		except:
			return GSI_GPR

	# 稳定性
	def get_GSI_STAB(self, para):
		keyword = "稳定性分析："
		start = para.find(keyword) + len(keyword)
		end = para.find("。", start)
		GSI_STAB = para[start: end]

		if GSI_STAB is not None:
			GSI_STAB = GSI_STAB.split("，")[-1]
			key = "围岩"
			GSI_STAB = GSI_STAB[GSI_STAB.find(key) + len(key):]
		else:
			GSI_STAB = "无"
		return GSI_STAB

	# 设计围岩级别
	def get_GSI_DSCR(self, table):
		GSI_DSCR = table.rows[1].cells[3].text
		return GSI_DSCR

    # 预报围岩级别
	def get_GSI_PSRL(self, table):
		GSI_PSRL = table.rows[1].cells[-1].text
		return GSI_PSRL

	# 结构构造
	def get_GSI_STRU(self, para):
		keyword = "结构构造："
		start = para.find(keyword) + len(keyword)
		end = para.find("。", start)
		GSI_STRU = para[start: end]
		if GSI_STRU == "":
			GSI_STRU = "无"
		return GSI_STRU


	# 获取附录表格
	def get_appendix(self, file):
		with plb.open(file) as pdf:
			page = pdf.pages[-1]
			table = page.extract_tables()[0]

			for row in table:
				while None in row:
					row.remove(None)

		return table

    # 岩性
	def get_GSI_LITH(self, para):
		keyword = "岩性："
		start = para.find(keyword) + len(keyword)
		end = para.find("。", start)
		GSI_LITH = para[start: end]

		keyword = "风化"
		GSI_LITH = GSI_LITH[GSI_LITH.find(keyword) + len(keyword):]
		if GSI_LITH is None:
			GSI_LITH = "无"
		return GSI_LITH

    # 风化程度
	def get_GSI_WEA(self, para):
		keyword = "岩性："
		start = para.find(keyword) + len(keyword)
		end = para.find("。", start)
		GSI_WEA = para[start: end]

		keyword = "风化"
		GSI_WEA = GSI_WEA[:GSI_WEA.find(keyword) + len(keyword)]
		if GSI_WEA is None:
			GSI_WEA = "无"
		return GSI_WEA

	# 地下水对应等级
	def get_GSI_WATG(self, table):
		return "无"

class PicturePDF:
    def __init__(self, type_name, file_name, input_path):
        self.file = file_name
        self.directory = self.parse_file(type_name, file_name)
        self.pixes = self.extract_graphs(input_path)

        # 未实现图片筛选
    def extract_graphs(self, input_path):
        pixes = []
        pdf = fitz.open(input_path)

        # 使用正则表达式来查找图片
        checkXO = r"/Type(?= */XObject)"
        checkIM = r"/Subtype(?= */Image)"

        # 获取对象数量长度
        lenXREF = pdf._getXrefLength()

        # 遍历每一个对象
        for i in range(1, lenXREF):
            # 定义对象字符串
            text = pdf._getXrefString(i)

            # 判断是否为对象或图片，若均不是则跳过
            isXObject = re.search(checkXO, text)
            isImage = re.search(checkIM, text)
            
            if not isXObject or not isImage:
                continue

            # 根据索引生成图像对象
            pix = fitz.Pixmap(pdf, i)
            pixes.append(pix)

        return pixes

    def parse_file(self, type_name, file_name):
        stage = None
        match = re.search("\d{3}", file_name)
        if match is not None:
            span = match.span()
            stage = file_name[span[0]: span[1]]
            stage = str(int(stage))

        GSI_INTE = None
        match = re.search("K\d\+\d{3}[-~](K\d\+)?\d{3}", file_name)
        if match is not None:
            span = match.span()
            GSI_INTE = file_name[span[0]: span[1]]
            if "-" in GSI_INTE:
                GSI_INTE = GSI_INTE.split("-")
                pre = GSI_INTE[0][: 3]
                GSI_INTE[1] = pre + GSI_INTE[1]
                GSI_INTE = "~".join(GSI_INTE)

        prefix = util.map_prefix(util.parse_prefix(file_name))

        return type_name + prefix + stage + "期" + GSI_INTE



class Processor(FileProcessBasic):
    name = "GPR-S3S4标"

    def save_fig(self, base, pictures, docx):
        base = os.path.join(base, "图片数据")
        util.checkout_directory(base)
        pic_dir = os.path.join(base, pictures.directory)
        util.checkout_directory(pic_dir)
        processed_pics = set()
        for i, p_id in enumerate(pictures.picture_ids):
            if not processed_pics.__contains__(p_id):
                processed_pics.add(p_id)
            else:
                continue
            img = docx.part.related_parts[p_id]
            file_type = img.filename.split(".")[-1]
            with open(os.path.join(pic_dir, "{}.{}".format(str(i + 1), file_type)), "wb") as f:
                f.write(img.blob)

    def save(self, output, record):
        output_path = os.path.join(output, "GPR_S3S4.csv")
        header = record.dict.keys()
        check_output_file(output_path, header)

        with open(output_path, "a+", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, record.dict.keys())
            w.writerow(record.dict)

    def run(self, input_path, output_path):
        files_to_process = set()
        files_to_transform = set()
        pdf_to_process = set()
        for file in os.listdir(input_path):
            absolute_file_path = os.path.join(input_path, file)
            if file.endswith(".doc"):
                files_to_transform.add(absolute_file_path)
            elif file.endswith(".docx"):
                files_to_process.add(absolute_file_path)
            elif file.endswith(".pdf"):
                pdf_to_process.add(absolute_file_path)
        files_to_delete = util.batch_doc_to_docx(files_to_transform)
        files_to_process = files_to_process.union(files_to_delete)

        for file in files_to_process:
            docx = Document(file)
            record = Record(docx)
            self.save(output_path, record)

            pics = Picture(Processor.name, file.split("\\")[-1], docx)
            self.save_fig(output_path, pics, docx)

            print("提取完成" + file)
        
        for file in pdf_to_process:
            record = RecordPDF(file)
            self.save(output_path, record)
            print("提取完成" + file)

        for file in files_to_delete:
            if os.path.exists(file):
                os.remove(file)

if __name__ == "__main__":
    test = Processor()
    inputpath = "/Users/budi/Desktop/iS3/Word"
    outputpath = "/Users/budi/Desktop/iS3/Word"
    test.run(inputpath, outputpath)