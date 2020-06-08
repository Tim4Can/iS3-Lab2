from docx import Document
import os
import csv
import re
from library.FileProcessBasic import FileProcessBasic
import util
import fitz
import xml.etree.cElementTree as ET
import pdfplumber as plb

class Record:

    def __init__(self, inte):
        # 稳定性
        GSI_STAB = ""
        # 设计围岩级别
        GSI_DSCR = ""
        # 推测围岩级别
        GSI_ESRG = ""
        #
        # GSI_STRU = self.get_GSI_STRU(para_suggestion)
        #
        # 桩号区间
        GSI_INTE = inte
        # 岩性
        GSI_LITH = ""
        # 风化程度
        GSI_WEA = ""
        # GSI_FAUL = self.get_GSI_FAUL(appendix)
        # GSI_FAUL = ""
        # GSI_WATG = self.get_GSI_WATG(appendix)
        # 密度
        GSI_DENST = ""
        # 纵波速度
        GSI_PWL = ""
        # 横波速度
        GSI_SWL = ""
        #泊松比
        GSI_PR = ""
        # 动态杨氏模量
        GSI_DYM = ""
        # 预报结果描述
        GSI_RESULT = ""
        # 完整性
        GSI_ITGT = ""
        # 地下水
        GSI_WATER = ""
        # 特殊地质情况
        GSI_SGS = ""


        self.dict = {

            "桩号区间": GSI_INTE,
            "岩性": GSI_LITH,
            "风化程度": GSI_WEA,
            "密度": GSI_DENST,
            "纵波速度": GSI_PWL,
            "横波速度": GSI_SWL,
            "泊松比": GSI_PR,
            "动态杨氏模量": GSI_DYM,
            "预报结果描述": GSI_RESULT,
            "地下水": GSI_WATER,
            "特殊地质情况": GSI_SGS,
            "稳定性": GSI_STAB,
            "设计围岩级别": GSI_DSCR,
            "推测围岩级别": GSI_ESRG,
            "完整性": GSI_ITGT
        }
    def get_attribute(self, table, i):
        self.get_GSI_LITH(table,i)
        self.get_GSI_RESULT(table,i)
        self.get_GSI_STAB(table,i)
        self.get_GSI_ITGT(table,i)
        self.get_GSI_DSCR(table,i)
        self.get_GSI_PSRL(table,i)


    # 预报结果描述
    def get_GSI_RESULT(self, table, j):
        for col in table.columns:
            if col.cells[0].text.strip().replace("\n", "").replace(" ", "") == "围岩主要工程地质条件"  and  col.cells[1].text.strip().replace("\n", "").replace(" ", "") == "主要工程地质特征":
                cvalue = col.cells[j].text.strip()
                # print(cvalue)
                keywords = "风化"
                prewords = "岩"
                pre = cvalue.find(prewords) + len(prewords)
                start = cvalue.find("，", pre) + len("，")
                end = cvalue.find("风化", start)
                # print(end)
                if end > 0:
                    self.dict["风化程度"] = cvalue[start:end + len(keywords)].replace("\n","")
                else:
                    self.dict["风化程度"] = "无"
                water_keywords = "含"
                water_start = cvalue.find(water_keywords)
                if not water_start < 0:
                    cvalue.replace("；", "，").replace("。", "，").replace("：", "，")
                    water_end = cvalue.find("，", water_start)
                    self.dict["地下水"] = cvalue[water_start:water_end].replace("\n","")
                else:
                    self.dict["地下水"] = "无"

                self.dict["预报结果描述"] = cvalue.replace("\n","")
                if self.dict["预报结果描述"] == "":
                    self.dict["预报结果描述"] = "无"
                break


    # 岩性
    def get_GSI_LITH(self, table, j):
        for col in table.columns:
            cvalue = col.cells[0].text.strip().replace("\n", "").replace(" ", "")
            if cvalue == "岩性":
                cvalue = col.cells[j].text.strip().replace("\n", "").replace(" ", "")
                # print(self.SB)
                # print(cvalue)
                if not cvalue == "":
                    self.dict["岩性"] = cvalue.replace("\n","")
                else:
                    self.dict["岩性"] = "无"

                break


    # 稳定性
    def get_GSI_STAB(self, table, j):
        for col in table.columns:
            cvalue = col.cells[0].text.strip().replace("\n", "").replace(" ", "")
            if cvalue == "围岩开挖后的稳定状态":
                cvalue = col.cells[j].text.strip().replace("\n", "").replace(" ", "")
                # print(self.SB)
                # print(cvalue)
                if not cvalue == "":
                    self.dict["稳定性"] = cvalue.replace("\n","")
                else:
                    self.dict["稳定性"] = "无"

                break

    # 完整性
    def get_GSI_ITGT(self,table, j):
        for col in table.columns:
            cvalue = col.cells[1].text.strip().replace("\n", "").replace(" ", "")
            if cvalue == "结构特征和完整状态":
                cvalue = col.cells[j].text.strip().replace("\n", "").replace(" ", "")
                if not cvalue == "":
                    self.dict["完整性"] = cvalue.replace("\n","")
                else:
                    self.dict["完整性"] = "无"

                break

    # 设计围岩级别
    def get_GSI_DSCR(self, table, j):
        for col in table.columns:
            cvalue = col.cells[0].text.strip().replace("\n", "").replace(" ", "")
            if cvalue == "设计围岩级别":
                cvalue = col.cells[j].text.strip().replace("\n", "").replace(" ", "")
                # print(self.SB)
                # print(cvalue)
                if cvalue != "" and cvalue != "/":
                    self.dict["设计围岩级别"] = cvalue[0:cvalue.find("级")].replace("\n","")
                else:
                    self.dict["设计围岩级别"] = "无"

                break

    # 推测围岩级别
    def get_GSI_PSRL(self, table, j):
        for col in table.columns:
            cvalue = col.cells[0].text.strip().replace("\n", "").replace(" ", "")
            if cvalue == "预报围岩级别":
                cvalue = col.cells[j].text.strip().replace("\n", "").replace(" ", "")
                # print(self.SB)
                # print(cvalue)
                if cvalue != "" and cvalue != "/":
                    self.dict["推测围岩级别"] = cvalue[0:cvalue.find("级")].replace("\n","")
                else:
                    self.dict["推测围岩级别"] = "无"

                break

    # # 地下水状态描述
    # def get_GSI_WATE(self):
    #     return "无"

class Picture:
    def __init__(self, type_name, file_name, docx):
        self.file = file_name
        self.directory = self.parse_file(type_name, file_name)
        self.picture_ids = self.extract_graphs(docx)

    def extract_graphs(self, docx):
        ids = []
        flag = False
        for i, p in enumerate(docx.paragraphs):
            if not flag and p.text.replace(" ", "").strip() == "目录":
                flag = True
            if flag:
                root = ET.fromstring(p._p.xml)
                pic_str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
                pics = root.findall(pic_str)
                image_str = "*/{urn:schemas-microsoft-com:vml}shape/{urn:schemas-microsoft-com:vml}imagedata"
                for pic in pics:
                    pict = pic.findall(image_str)
                    if len(pict) > 0:
                        text = docx.paragraphs[i + 1].text
                        if not text.endswith("示意图"):
                            ids.append(pict[0].attrib[
                                           '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'])
        return ids

    def parse_file(self, type_name, file_name):
        stage = None
        match = re.search("\d{3}", file_name)
        if match is not None:
            span = match.span()
            stage = file_name[span[0]: span[1]]
            stage = str(int(stage))

        SKTH_INTE = None
        match = re.search("K\d\+\d{3}[-~](K\d\+)?\d{3}", file_name)
        if match is not None:
            span = match.span()
            SKTH_INTE = file_name[span[0]: span[1]]
            if "-" in SKTH_INTE:
                SKTH_INTE = SKTH_INTE.split("-")
                pre = SKTH_INTE[0][: 3]
                SKTH_INTE[1] = pre + SKTH_INTE[1]
                SKTH_INTE = "~".join(SKTH_INTE)
        else:
            match=re.search("LS\d\+\d{3}[-~](LS\d\+)?\d{3}",file_name)
            if match is not None:
                span=match.span()
                SKTH_INTE = file_name[span[0]: span[1]]
                if "-" in SKTH_INTE:
                    SKTH_INTE = SKTH_INTE.split("-")
                    pre = SKTH_INTE[0][: 3]
                    SKTH_INTE[1] = pre + SKTH_INTE[1]
                    SKTH_INTE = "~".join(SKTH_INTE)
            else:
                SKTH_INTE = ""




        prefix = util.map_prefix(util.parse_prefix(file_name))

        return type_name + prefix + stage + "期" + SKTH_INTE


class PicturePDF:
    def __init__(self, type_name, file_name, input_path):
        self.file = file_name
        self.directory = self.parse_file(type_name, file_name)
        self.pixes = self.extract_graphs(input_path)

        # 未实现图片筛选

    def extract_graphs(self, input_path):
        pixes = []
        pdf = fitz.open(input_path)

        # 使用正则表达式来查找图片
        checkXO = r"/Type(?= */XObject)"
        checkIM = r"/Subtype(?= */Image)"

        # 获取对象数量长度
        lenXREF = pdf._getXrefLength()

        # 遍历每一个对象
        for i in range(1, lenXREF):
            # 定义对象字符串
            text = pdf._getXrefString(i)

            # 判断是否为对象或图片，若均不是则跳过
            isXObject = re.search(checkXO, text)
            isImage = re.search(checkIM, text)

            if not isXObject or not isImage:
                continue

            # 根据索引生成图像对象
            pix = fitz.Pixmap(pdf, i)
            if pix.w > 180 and pix.h > 150:
                pixes.append(pix)

        # titles = []
        # with plb.open(input_path) as pdf_text:
        #     texts = [pdf_text.pages[i].extract_text() for i in range(len(pdf_text.pages))]
        #     for text in texts:
        #         pattern = r"^\s*图\s*\d.*\n"
        #         result = re.findall(pattern, text, re.M)
        #         titles.extend(result)
        # filtered_pics = []
        # if len(titles) == len(pixes):
        #     for i, title in enumerate(titles):
        #         title = title.replace("\n", "").strip()
        #         if not title.endswith("示意图"):
        #             filtered_pics.append(pixes[i])
        # else:
        #     filtered_pics = pixes
        # return filtered_pics
        return pixes

    def parse_file(self, type_name, file_name):
        stage = None
        match = re.search("\d{3}", file_name)
        if match is not None:
            span = match.span()
            stage = file_name[span[0]: span[1]]
            stage = str(int(stage))

        GSI_INTE = None
        match = re.search("K\d\+\d{3}[-~](K\d\+)?\d{3}", file_name)
        if match is not None:
            span = match.span()
            GSI_INTE = file_name[span[0]: span[1]]
            if "-" in GSI_INTE:
                GSI_INTE = GSI_INTE.split("-")
                pre = GSI_INTE[0][: 3]
                GSI_INTE[1] = pre + GSI_INTE[1]
                GSI_INTE = "~".join(GSI_INTE)
        else:
            match=re.search("LS1 \d\+\d{3}[-~](LS1 \d\+)?\d{3}",file_name)
            if match is not None:
                span = match.span()
                GSI_INTE = file_name[span[0]: span[1]]
                if "-" in GSI_INTE:
                    GSI_INTE = GSI_INTE.split("-")
                    pre = GSI_INTE[0][: 3]
                    GSI_INTE[1] = pre + GSI_INTE[1]
                    GSI_INTE = "~".join(GSI_INTE)
            else:
                GSI_INTE=""

        prefix = util.map_prefix(util.parse_prefix(file_name))

        return type_name + prefix + stage + "期" + GSI_INTE

class Processor(FileProcessBasic):
    name = "TSP-S3S4标"

    def save(self, output, records):
        output_path = os.path.join(output, "TSP_S3S4.csv")
        for record in records:
            header = record.dict.keys()
            util.check_output_file(output_path, header)

            with open(output_path, "a+", encoding="utf_8_sig", newline="") as f:
                w = csv.DictWriter(f, record.dict.keys())
                w.writerow(record.dict)

    def save_fig(self, base, pictures, docx):
        base = os.path.join(base, "图片数据")
        util.checkout_directory(base)
        pic_dir = os.path.join(base, pictures.directory)
        util.checkout_directory(pic_dir)
        processed_pics = set()
        for i, p_id in enumerate(pictures.picture_ids):
            if not processed_pics.__contains__(p_id):
                processed_pics.add(p_id)
            else:
                continue
            img = docx.part.related_parts[p_id]
            file_type = img.filename.split(".")[-1]
            with open(os.path.join(pic_dir, "{}.{}".format(str(i + 1), file_type)), "wb") as f:
                f.write(img.blob)

    def save_fig_PDF(self, base, pictures):
        base = os.path.join(base, "图片数据")
        util.checkout_directory(base)
        pic_dir = os.path.join(base, pictures.directory)
        util.checkout_directory(pic_dir)
        for i, pix in enumerate(pictures.pixes):
            new_name = "{}.png".format(i + 1)
            # 如果pix.n<5,可以直接存为PNG
            if pix.n < 5:
                path = os.path.join(pic_dir, new_name)
                pix.writePNG(path)
            # 否则先转换CMYK
            else:
                pix0 = fitz.Pixmap(fitz.csRGB, pix)
                pix0.writePNG(os.path.join(pic_dir, new_name))
                pix0 = None
            # 释放资源
            pix = None


    def run(self, input_path, output_path):
        files_to_process = set()
        files_to_transform = set()
        pdf_to_process=set()
        for file in os.listdir(input_path):
            absolute_file_path = os.path.join(input_path, file)
            if file.endswith(".doc"):
                files_to_transform.add(absolute_file_path)
            elif file.endswith(".docx"):
                files_to_process.add(absolute_file_path)
            elif file.endswith(".pdf"):
                pdf_to_process.add(absolute_file_path)
        files_to_delete = util.batch_doc_to_docx(files_to_transform)
        files_to_process = files_to_process.union(files_to_delete)

        for file in files_to_process:
            docx = Document(file)
            records = list()
            table2 = docx.tables[-2]
            table3 = docx.tables[-1]
            conclusion = self.get_conclusion(docx)

            self.get_record_table3(records,table3)
            self.get_record_table2(records,table2)
            self.get_record_conclusion(records, conclusion)
            # 图片提取
            pics = Picture(Processor.name, file.split("\\")[-1], docx)
            self.save_fig(output_path, pics, docx)

            self.save(output_path, records)
            print("提取完成" + file)

        for file in pdf_to_process:
            docx = Document()
            with plb.open(file) as pdf:
                tables = []
                content=""
                for i in range(len(pdf.pages)):
                    table=pdf.pages[i].extract_tables()
                    content += pdf.pages[i].extract_text()

                    if not len(table)==0:
                        tables.append(table)
                tb2 = tables[-2]
                table2 = docx.add_table(len(tb2[0]), len(tb2[0][0]))
                table2 = self.traverse_table(tb2, table2)
                tb3 = tables[-1]
                table3 = docx.add_table(len(tb3[0]),len(tb3[0][0]))
                table3 = self.traverse_table(tb3, table3)
                records = list()
                conclusion = self.get_pdf_conclusion(content)

                self.get_record_table3(records, table3)
                self.get_record_table2(records, table2)
                self.get_record_conclusion(records, conclusion)
            self.save(output_path, records)

            # 提取PDF图片
            pics_PDF = PicturePDF(Processor.name, file.split("\\")[-1], file)
            self.save_fig_PDF(output_path, pics_PDF)
            print("提取完成" + file)



        for file in files_to_delete:
            if os.path.exists(file):
                os.remove(file)

    def get_pdf_conclusion(self, content):
        para_conclusion = ""
        flag = 0
        lines = content.splitlines()
        for line in lines:
            line = line.strip()
            # 无意义，略过
            if line == "" or (line.startswith("第") and line.endswith("页")):
                continue

            # 提取探测结果
            if line.startswith("8"):
                flag = 1
                continue
            if flag == 1:
                para_conclusion += line

        return para_conclusion

    def get_conclusion(self, docx):
        para_conclusion = ""
        flag = 0
        for i, p in enumerate(docx.paragraphs):
            if p.text.startswith("8"):
                flag = 1
                continue
            if flag == 1:
                para_conclusion += p.text

        return para_conclusion

    def get_record_conclusion(self, records, conclusion):
        conclusion = conclusion.replace("施工建议：", "（")
        for record in records:
            keywords = conclusion.find(record.dict["桩号区间"])
            if not keywords < 0:
                start = conclusion.find("段：", keywords)
                if not start < 0:
                    end = conclusion.find("（", start)
                    if end < 0:
                        end = len(conclusion)
                    record.dict["特殊地质情况"] = conclusion[start + len("段："):end].replace("\n","")
                else:
                    end = conclusion.find("（", keywords)
                    if end < 0:
                        end = len(conclusion)
                    record.dict["特殊地质情况"] = conclusion[keywords + len(record.dict["桩号区间"]):end].replace("\n","")
                conclusion = conclusion[end:len(conclusion)]
            else:
                record.dict["特殊地质情况"] = "无"

    def get_record_table2(self, records, table2):
        mileages, vps, vms, prs, dss = self.get_info_table2(table2)
        for record in records:
            inte = record.dict["桩号区间"]
            # print(inte)
            # print("inte:")
            # print(inte)
            # match=re.search("[K]")
            start = inte.find("K")
            if start < 0:
                start = inte.find("LS")
                if start >= 0:
                    start = start + len("LS")
                    # print(start)
            else:
                start = start + len("K")


            end = inte.find("～",start)
            # print(inte[start:end].replace("+",""))
            max = int(inte[start:end].replace("+",""))
            # print("max:")
            # print(max)
            start = inte.find("K",end)
            if start < 0:
                start = inte.find("LS",end)
                if start >= 0:
                    start = start + len("LS")
            else:
                start = start + len("K")

            min = int(inte[start:len(inte)].replace("+",""))
            tvps = []
            tvms = []
            tprs = []
            tdss = []
            for i in range(0, len(mileages)):

                mileage = mileages[i]

                temp = int(mileage.replace(",","").replace("，",""))
                if temp >= min and temp <= max:
                    tvps.append(vps[i])
                    tvms.append(vms[i])
                    tprs.append(prs[i])
                    tdss.append(dss[i])
            tvps.sort()
            tvms.sort()
            tprs.sort()
            tdss.sort()
            if not len(tvps) == 0:
                record.dict["纵波速度"] = tvps[0] + "~" + tvps[len(tvps) - 1]
            else:
                record.dict["纵波速度"] = "无"
            if not len(tvms) == 0:
                record.dict["横波速度"] = str(tvms[0]) + "~" + str(tvms[len(tvms) - 1])
            else:
                record.dict["横波速度"] = "无"
            if not len(tprs) == 0:
                record.dict["泊松比"] = str(tprs[0]) + "~" + str(tprs[len(tprs) - 1])
            else:
                record.dict["泊松比"] = "无"
            if not len(tdss) == 0:
                record.dict["密度"] = str(tdss[0]) + "~" + str(tdss[len(tdss) - 1])
            else:
                record.dict["密度"] = "无"

                # print(record.dict["纵波速度"])



    def get_info_table2(self, table):
        i = ma = vp = vpm = pr = ds = 0
        row = table.rows[0]
        while i < len(table.columns):
            if row.cells[i].text.strip().replace("\n", "").replace(" ", "") == "里程":
                ma = i
            elif row.cells[i].text.strip().replace("\n", "").replace(" ", "") == "Vp(m/s)":
                vp = i
            elif row.cells[i].text.strip().replace("\n", "").replace(" ", "") == "Vp/Vs":
                vpm = i
            elif row.cells[i].text.strip().replace("\n", "").replace(" ", "") == "泊松比":
                pr = i
            elif row.cells[i].text.strip().replace("\n", "").replace(" ", "") == "密度(g/cm3)":
                ds = i
            i = i + 1
        mileages = []
        vps = []
        vms = []
        prs = []
        dss = []
        for j in range(2, len(table.rows)):
            row = table.rows[j]
            mileages.append(row.cells[ma].text.strip().replace("-",""))
            # print("mileage:"+row.cells[ma].text.strip().replace("-",""))
            vps.append(row.cells[vp].text.strip().replace(",","").replace("，",""))
            vm = round(float(row.cells[vp].text.strip().replace(",","").replace("，","")) / float(row.cells[vpm].text.strip()),2)
            vms.append(vm)
            prs.append(float(row.cells[pr].text.strip()))
            dss.append(float(row.cells[ds].text.strip()))
        return mileages, vps, vms, prs, dss

    def traverse_table(self, tb, docx_table):
        table=tb[0]
        flag = 0
        for i in range(len(table)):
            row = docx_table.rows[i]
            cols=table[i]

            for k in range(len(table[i])):

                if not cols[k]==None:
                    if cols[k].replace("\n","" ).replace(" ","")=="预报里程范围":
                        flag=1
                    if flag == 1 and k == 3 and i != 0 and i != 1:
                        strs = cols[k].split("\n")
                        for i, s in enumerate(strs):
                            if s == "":
                                strs.remove(s)
                                continue
                            strs[i] = s.split(" ")

                        row_num = len(strs[0])

                        result = ""
                        for i in range(row_num):
                            for j in range(len(strs)):
                                result += strs[j][row_num - i - 1]
                        cols[k] = result

                    content=cols[k]
                    row.cells[k].text = content
                else:
                    if i > 0:
                        table[i][k] = table[i-1][k]
                    elif k > 0:
                        table[i][k] = table[i][k-1]
                    row.cells[k].text = table[i][k]



        return docx_table




    def get_record_table3(self, records, table):
        for col in table.columns:
            cvalue = col.cells[0].text.strip().replace("\n", "").replace(" ", "")
            if cvalue == "预报里程范围":
                # print("预报里程范围")
                for i in range(2, len(col.cells)):
                    cvalue = col.cells[i].text.strip().replace("\n", "").replace(" ", "")
                    # print(cvalue)
                    record = Record(cvalue)
                    record.get_attribute(table, i)
                    records.append(record)
                break



if __name__ == "__main__":
    test = Processor()
    inputpath = "e:/study/is3/tsp2"
    outputpath = "C:/Users/DELL/Desktop/iS3"
    test.run(inputpath, outputpath)
