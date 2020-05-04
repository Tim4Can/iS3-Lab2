from docx import Document
import os
import csv
from library.FileProcessBasic import FileProcessBasic
import util
import re
import fitz
import time
import pdfplumber as plb
import pandas as pd
import xml.etree.cElementTree as ET


# word:桩号区间list
def get_INTE(table):
    INTE = []
    for i in range(len(table.rows)):
        if table.cell(i, 1).text.strip().replace("\n", "") == "里程段（m）":
            for j in range(len(table.rows) - 1):
                per_INTE = table.cell(i + j + 1, 1).text.strip().replace("\n", "")
                first_CHAI = per_INTE.split("～")[0]
                first_CHAI_prefix = first_CHAI.split("+")[0]
                first_others = first_CHAI.split("+")[1]
                second_CHAI = per_INTE.split("～")[1]
                second_CHAI_prefix = second_CHAI.split("+")[0]
                second_others = second_CHAI.split("+")[1]
                if per_INTE == "":
                    per_INTE = "无"
                if second_CHAI_prefix == "":
                    first = (first_CHAI_prefix, first_others)
                    second = (first_CHAI_prefix, second_others)
                    first_CHAI = "+".join(first)
                    second_CHAI = "+".join(second)
                    per_INTE = (first_CHAI, second_CHAI)
                    per_INTE = "～".join(per_INTE)

                INTE.append(per_INTE)
    return INTE


# pdf:桩号区间list
def get_INTE_PDF(table):
    INTE = []
    for i in range(len(table)):
        if table[i][1].strip().replace("\n", "") == "里程段（m）":
            for j in range(len(table) - 1):
                per_INTE = table[i + j + 1][1].strip().replace("\n", "")
                first_CHAI = per_INTE.split("～")[0]
                first_CHAI_prefix = first_CHAI.split("+")[0]
                first_others = first_CHAI.split("+")[1]
                second_CHAI = per_INTE.split("～")[1]
                second_CHAI_prefix = second_CHAI.split("+")[0]
                second_others = second_CHAI.split("+")[1]
                if per_INTE == "":
                    per_INTE = "无"
                if second_CHAI_prefix == "":
                    first = (first_CHAI_prefix, first_others)
                    second = (first_CHAI_prefix, second_others)
                    first_CHAI = "+".join(first)
                    second_CHAI = "+".join(second)
                    per_INTE = (first_CHAI, second_CHAI)
                    per_INTE = "～".join(per_INTE)

                INTE.append(per_INTE)
    return INTE


class Record:
    def __init__(self, docx, INTE):
        table_result = docx.tables[1]
        table_analysis = docx.tables[2]

        name = self.get_cover(docx)
        para_conclusion, para_suggestion = self.locate_paragraph(docx)

        GPRF_INTE = util.parse_GPRF_INTE(name, INTE)

        GPRF_MD, GPRF_ZPS, GPRF_HPS, GPRF_PSB, GPRF_EM = self.get_GPRF_table_result(table_result, INTE)
        GPRF_FORE, GPRF_WEA, GPRF_WATE, GPRF_INTE2, GPRF_STAB, GPRF_PSRL = self.get_GPRF_table_analysis(
            table_analysis,
            INTE)
        GPRF_LITH, GPRF_FAUL = self.get_GPRF_LITH_and_GPRF_FAUL(para_conclusion, INTE)

        # 未实现
        GPRF_STRE = self.get_GPRF_STRE()

        self.dict = {
            "桩号区间": GPRF_INTE,
            "岩性": GPRF_LITH,
            "密度ρ（g/cm3）": GPRF_MD,
            "纵波速度Vp（m/s）": GPRF_ZPS,
            "横波速度Vs（m/s）": GPRF_HPS,
            "泊松比σ": GPRF_PSB,
            "动态杨氏模量（GPa）": GPRF_EM,
            "预报结果描述": GPRF_FORE,
            "风化程度": GPRF_WEA,
            "地下水": GPRF_WATE,
            "完整性": GPRF_INTE2,
            "稳定性": GPRF_STAB,
            "特殊地质情况": GPRF_FAUL,
            "推测围岩级别": GPRF_PSRL,
            "设计围岩级别": GPRF_STRE,

        }

    def get_cover(self, docx):
        name = None
        for paragraph in docx.paragraphs:
            if paragraph.text.startswith("隧道名称：") or paragraph.text.startswith("项目名称："):
                name = paragraph.text.split("：")[1].strip()
            if name is not None:
                return name

    def locate_paragraph(self, docx):
        para_conclusion = ""  # 7.1 结论
        para_suggestion = ""  # 7.2 建议
        for i, p in enumerate(docx.paragraphs):
            if p.text.startswith("7.1"):
                i += 1
                p = docx.paragraphs[i]
                while not p.text.startswith("7.2"):
                    para_conclusion += p.text
                    i += 1
                    p = docx.paragraphs[i]
                i -= 1
            elif p.text.startswith("7.2"):
                i += 1
                p = docx.paragraphs[i]
                while not p.text.startswith("8") or p.text == "\n":
                    para_suggestion += p.text
                    i += 1
                    p = docx.paragraphs[i]
        return para_conclusion, para_suggestion

    # 表6.2 中 预报结果/预报推断结果 单元格
    def get_GPRF_table_analysis(self, table, val):
        GPRF_FORE = ""  # 预报结果/预报推断结果
        GPRF_WEA = ""  # 风化程度
        GPRF_WATE = ""  # 地下水
        GPRF_INTE2 = ""  # 完整性
        GPRF_STAB = ""  # 稳定性
        GPRF_PSRL = ""  # 推测围岩级别/推断围岩级别
        for i in range(len(table.rows)):
            text = table.cell(i, 0).text
            text = text.strip().split('\n')[:-1]
            text = "".join(text)
            if text != "" and (val in text or text in val):
                for j in range(5):
                    if table.cell(3, j).text.strip() == "预报结果" or "预报推断结果":
                        GPRF_FORE_whole = table.cell(i, j).text.strip().replace("\n", "")

                        # 预报结果/预报推断结果
                        fore_start = 0
                        fore_index = GPRF_FORE_whole.find("围岩")
                        fore_end = GPRF_FORE_whole.find("。", fore_index)
                        GPRF_FORE = GPRF_FORE_whole[fore_start:fore_end] + "。"
                        if GPRF_FORE == "":
                            GPRF_FORE = "无"

                        # 风化程度
                        wea_index = GPRF_FORE_whole.find("风化")
                        wea_start1 = GPRF_FORE_whole.rfind("，", 0, wea_index) + 1
                        wea_start2 = GPRF_FORE_whole.rfind("。", 0, wea_index) + 1
                        if wea_start1 < wea_start2:
                            wea_start = wea_start2
                        else:
                            wea_start = wea_start1
                        wea_end = GPRF_FORE_whole.find("化", wea_start) + 1
                        GPRF_WEA = GPRF_FORE_whole[wea_start: wea_end].replace("岩体", "")
                        if GPRF_WEA == "":
                            GPRF_WEA = "无"

                        # 地下水
                        wate_start = GPRF_FORE_whole.rfind("岩体含水")
                        wate_end = GPRF_FORE_whole.find("，", wate_start)
                        if wate_start != -1:
                            GPRF_WATE = GPRF_FORE_whole[wate_start: wate_end]
                        elif wate_start == -1:

                            wate_index = GPRF_FORE_whole.rfind("完整性")
                            wate_index1 = GPRF_FORE_whole.find("水", 0, wate_index)
                            if wate_index1 != -1:
                                wate_index2 = GPRF_FORE_whole.rfind("水", 0, wate_index)
                                wate_start1 = GPRF_FORE_whole.rfind("，", 0, wate_index1) + 1
                                wate_start2 = GPRF_FORE_whole.rfind("；", 0, wate_index1) + 1
                                wate_end = GPRF_FORE_whole.find("，", wate_index2, wate_index)
                                if wate_start1 < wate_start2:
                                    wate_start = wate_start2
                                else:
                                    wate_start = wate_start1
                                GPRF_WATE = GPRF_FORE_whole[wate_start: wate_end]
                            else:
                                GPRF_WATE = "无"

                        # 完整性
                        inte2_start = GPRF_FORE_whole.find("完整程度")
                        inte2_index = GPRF_FORE_whole.find("结构")
                        inte2_end = GPRF_FORE_whole.rfind("，", inte2_start, inte2_index)
                        GPRF_INTE2 = GPRF_FORE_whole[inte2_start: inte2_end]
                        if GPRF_INTE2 == "":
                            GPRF_INTE2 = "无"

                        # 稳定性
                        stab_start = GPRF_FORE_whole.find("稳定性")
                        stab_end = GPRF_FORE_whole.find("。", stab_start)
                        GPRF_STAB = GPRF_FORE_whole[stab_start: stab_end]
                        if GPRF_STAB == "":
                            GPRF_STAB = "无"

                        # # 特殊地质情况
                        # faul_start = GPRF_FORE_whole.find("裂隙")
                        # faul_end = GPRF_FORE_whole.find("育", faul_start) + 1
                        # GPRF_FAUL = GPRF_FORE_whole[faul_start: faul_end]
                        # if GPRF_FAUL == "":
                        #     GPRF_FAUL = "无"

                        # 推测围岩级别/推断围岩级别
                        GPRF_PSRL = table.cell(i, j + 1).text.strip().replace("\n", "")
                        if GPRF_PSRL == "":
                            GPRF_PSRL = "无"

        return GPRF_FORE, GPRF_WEA, GPRF_WATE, GPRF_INTE2, GPRF_STAB, GPRF_PSRL

    # 表6.1探测结果
    def get_GPRF_table_result(self, table, val):
        GPRF_MD = ""  # 密度
        GPRF_ZPS = ""  # 纵波速度
        GPRF_HPS = ""  # 横波速度
        GPRF_PSB = ""  # 泊松比
        GPRF_EM = ""  # 动态杨氏模量
        for i in range(len(table.rows) - 1):
            text = table.cell(i + 1, 1).text.strip().replace("\n", "")
            first_CHAI = text.split("～")[0]
            first_CHAI_prefix = first_CHAI.split("+")[0]
            first_others = first_CHAI.split("+")[1]
            second_CHAI = text.split("～")[1]
            second_CHAI_prefix = second_CHAI.split("+")[0]
            second_others = second_CHAI.split("+")[1]
            if second_CHAI_prefix == "":
                first = (first_CHAI_prefix, first_others)
                second = (first_CHAI_prefix, second_others)
                first_CHAI = "+".join(first)
                second_CHAI = "+".join(second)
                text = (first_CHAI, second_CHAI)
                text = "～".join(text)
            if text == val:
                # GPRF_LITH = table.cell(i + 1, 0).text
                GPRF_MD = table.cell(i + 1, 2).text
                GPRF_ZPS = table.cell(i + 1, 3).text
                GPRF_HPS = table.cell(i + 1, 4).text
                GPRF_PSB = table.cell(i + 1, 5).text
                GPRF_EM = table.cell(i + 1, 6).text

                # if GPRF_LITH == "":
                #     GPRF_LITH = "无"
                if GPRF_MD == "":
                    GPRF_MD = "无"
                if GPRF_ZPS == "":
                    GPRF_ZPS = "无"
                if GPRF_HPS == "":
                    GPRF_HPS = "无"
                if GPRF_PSB == "":
                    GPRF_PSB = "无"
                if GPRF_EM == "":
                    GPRF_EM = "无"
        return GPRF_MD, GPRF_ZPS, GPRF_HPS, GPRF_PSB, GPRF_EM

    def get_GPRF_LITH_and_GPRF_FAUL(self, para, val):
        GPRF_LITH = "无"  # 岩性
        GPRF_FAUL = "无"  # 特殊地质情况
        try:
            lith_start = para.find("岩性") + 2
            lith_end = para.find("，", lith_start)
            GPRF_LITH = para[lith_start:lith_end].replace("主要为", "")

            faul_index = para.rfind("围岩的")
            faul_start = para.find(val)
            if faul_index == -1:
                faul_end = para.rfind("。", faul_start)
            else:
                faul_end = faul_index
            GPRF_FAUL = para[faul_start:faul_end]
            if GPRF_FAUL == "":
                GPRF_FAUL = "无"

            return GPRF_LITH, GPRF_FAUL
        except:
            return GPRF_LITH, GPRF_FAUL

    # 设计围岩级别
    def get_GPRF_STRE(self):
        return "无"


class RecordPDF:
    def __init__(self, pdf, INTE, table_result, table_analysis, para_conclusion):

        name = self.get_cover(pdf)

        GPRF_INTE = util.parse_GPRF_INTE(name, INTE)
        GPRF_MD, GPRF_ZPS, GPRF_HPS, GPRF_PSB, GPRF_EM = self.get_GPRF_table_result(table_result, INTE)
        GPRF_FORE, GPRF_WEA, GPRF_WATE, GPRF_INTE2, GPRF_STAB, GPRF_PSRL = self.get_GPRF_table_analysis(
            table_analysis,
            INTE)
        GPRF_LITH, GPRF_FAUL = self.get_GPRF_LITH_and_GPRF_FAUL(para_conclusion, INTE)

        # 未实现
        GPRF_STRE = self.get_GPRF_STRE()

        self.dict = {
            "桩号区间": GPRF_INTE,
            "岩性": GPRF_LITH,
            "密度ρ（g/cm3）": GPRF_MD,
            "纵波速度Vp（m/s）": GPRF_ZPS,
            "横波速度Vs（m/s）": GPRF_HPS,
            "泊松比σ": GPRF_PSB,
            "动态杨氏模量（GPa）": GPRF_EM,
            "预报结果描述": GPRF_FORE,
            "风化程度": GPRF_WEA,
            "地下水": GPRF_WATE,
            "完整性": GPRF_INTE2,
            "稳定性": GPRF_STAB,
            "特殊地质情况": GPRF_FAUL,
            "推测围岩级别": GPRF_PSRL,
            "设计围岩级别": GPRF_STRE,

        }

    def get_cover(self, pdf):
        name = None
        contents = pdf.pages[0].extract_text().split("\n")
        for content in contents:
            if content.startswith("隧道名称：") or content.startswith("项目名称"):
                name = content.split("：")[1].strip()
            if name is not None:
                return name

    # 表6.2 中 预报结果/预报推断结果 单元格
    def get_GPRF_table_analysis(self, table, val):
        GPRF_FORE = ""  # 预报结果/预报推断结果
        GPRF_WEA = ""  # 风化程度
        GPRF_WATE = ""  # 地下水
        GPRF_INTE2 = ""  # 完整性
        GPRF_STAB = ""  # 稳定性
        GPRF_PSRL = ""  # 推测围岩级别/推断围岩级别
        for i in range(len(table)):
            text = str(table[i][0])
            text = text.strip().split('\n')[:-1]
            text = "".join(text)
            # print(table)
            if text != "" and (val in text or text in val):
                # for j in range(len(table[0])-1):
                if table[3][2].strip() == "预报结果" or "预报推断结果":
                    GPRF_FORE_whole = table[i][2].strip().replace("\n", "")

                    # 预报结果/预报推断结果
                    fore_start = 0
                    fore_index = GPRF_FORE_whole.find("围岩")
                    fore_end = GPRF_FORE_whole.find("。", fore_index)
                    GPRF_FORE = GPRF_FORE_whole[fore_start:fore_end] + "。"

                    if GPRF_FORE == "":
                        GPRF_FORE = "无"

                    # 风化程度
                    wea_index = GPRF_FORE_whole.find("风化")
                    wea_start1 = GPRF_FORE_whole.rfind("，", 0, wea_index) + 1
                    wea_start2 = GPRF_FORE_whole.rfind("。", 0, wea_index) + 1
                    if wea_start1 < wea_start2:
                        wea_start = wea_start2
                    else:
                        wea_start = wea_start1
                    wea_end = GPRF_FORE_whole.find("化", wea_start) + 1
                    GPRF_WEA = GPRF_FORE_whole[wea_start: wea_end].replace("岩体", "")
                    if GPRF_WEA == "":
                        GPRF_WEA = "无"

                    # 地下水
                    wate_start = GPRF_FORE_whole.rfind("岩体含水")
                    wate_end = GPRF_FORE_whole.find("，", wate_start)
                    if wate_start != -1:
                        GPRF_WATE = GPRF_FORE_whole[wate_start: wate_end]
                    elif wate_start == -1:

                        wate_index = GPRF_FORE_whole.rfind("完整性")
                        wate_index1 = GPRF_FORE_whole.find("水", 0, wate_index)
                        if wate_index1 != -1:
                            wate_index2 = GPRF_FORE_whole.rfind("水", 0, wate_index)
                            wate_start1 = GPRF_FORE_whole.rfind("，", 0, wate_index1) + 1
                            wate_start2 = GPRF_FORE_whole.rfind("；", 0, wate_index1) + 1
                            wate_end = GPRF_FORE_whole.find("，", wate_index2, wate_index)
                            if wate_start1 < wate_start2:
                                wate_start = wate_start2
                            else:
                                wate_start = wate_start1
                            GPRF_WATE = GPRF_FORE_whole[wate_start: wate_end]
                        else:
                            GPRF_WATE = "无"

                    # 完整性
                    inte2_start = GPRF_FORE_whole.find("完整程度")
                    inte2_index = GPRF_FORE_whole.find("结构")
                    inte2_end = GPRF_FORE_whole.rfind("，", inte2_start, inte2_index)
                    GPRF_INTE2 = GPRF_FORE_whole[inte2_start: inte2_end]
                    if GPRF_INTE2 == "":
                        GPRF_INTE2 = "无"

                    # 稳定性
                    stab_start = GPRF_FORE_whole.find("稳定性")
                    stab_end = GPRF_FORE_whole.find("。", stab_start)
                    GPRF_STAB = GPRF_FORE_whole[stab_start: stab_end]
                    if GPRF_STAB == "":
                        GPRF_STAB = "无"

                    # # 特殊地质情况
                    # faul_start = GPRF_FORE_whole.find("裂隙")
                    # faul_end = GPRF_FORE_whole.find("育", faul_start) + 1
                    # GPRF_FAUL = GPRF_FORE_whole[faul_start: faul_end]
                    # if GPRF_FAUL == "":
                    #     GPRF_FAUL = "无"

                    # 推测围岩级别/推断围岩级别
                    GPRF_PSRL = str(table[i][-1]).replace("\n", "")
                    if GPRF_PSRL == "":
                        GPRF_PSRL = "无"

        return GPRF_FORE, GPRF_WEA, GPRF_WATE, GPRF_INTE2, GPRF_STAB, GPRF_PSRL

    # 表6.1探测结果
    def get_GPRF_table_result(self, table, val):
        GPRF_MD = ""  # 密度
        GPRF_ZPS = ""  # 纵波速度
        GPRF_HPS = ""  # 横波速度
        GPRF_PSB = ""  # 泊松比
        GPRF_EM = ""  # 动态杨氏模量
        for i in range(len(table) - 1):
            text = table[i + 1][1].strip().replace("\n", "")
            first_CHAI = text.split("～")[0]
            first_CHAI_prefix = first_CHAI.split("+")[0]
            first_others = first_CHAI.split("+")[1]
            second_CHAI = text.split("～")[1]
            second_CHAI_prefix = second_CHAI.split("+")[0]
            second_others = second_CHAI.split("+")[1]
            if second_CHAI_prefix == "":
                first = (first_CHAI_prefix, first_others)
                second = (first_CHAI_prefix, second_others)
                first_CHAI = "+".join(first)
                second_CHAI = "+".join(second)
                text = (first_CHAI, second_CHAI)
                text = "～".join(text)
            if text == val:
                GPRF_MD = table[i + 1][2]
                GPRF_ZPS = table[i + 1][3]
                GPRF_HPS = table[i + 1][4]
                GPRF_PSB = table[i + 1][5]
                GPRF_EM = table[i + 1][6]

                # if GPRF_LITH == "":
                #     GPRF_LITH = "无"
                if GPRF_MD == "":
                    GPRF_MD = "无"
                if GPRF_ZPS == "":
                    GPRF_ZPS = "无"
                if GPRF_HPS == "":
                    GPRF_HPS = "无"
                if GPRF_PSB == "":
                    GPRF_PSB = "无"
                if GPRF_EM == "":
                    GPRF_EM = "无"
        return GPRF_MD, GPRF_ZPS, GPRF_HPS, GPRF_PSB, GPRF_EM

    def get_GPRF_LITH_and_GPRF_FAUL(self, para, val):
        GPRF_LITH = "无"  # 岩性
        GPRF_FAUL = "无"  # 特殊地质情况
        try:
            lith_start = para.find("岩性") + 2
            lith_end = para.find("，", lith_start)
            GPRF_LITH = para[lith_start:lith_end].replace("主要为", "")

            faul_index = para.rfind("围岩的")
            faul_start = para.find(val)
            if faul_index == -1:
                faul_end = para.rfind("。", faul_start)
            else:
                faul_end = faul_index
            GPRF_FAUL = para[faul_start:faul_end]
            if GPRF_FAUL == "":
                GPRF_FAUL = "无"

            return GPRF_LITH, GPRF_FAUL
        except:
            return GPRF_LITH, GPRF_FAUL

    # 设计围岩级别
    def get_GPRF_STRE(self):
        return "无"


class Picture:
    def __init__(self, type_name, file_name, docx):
        self.file = file_name
        self.directory = self.parse_file(type_name, file_name)
        self.picture_ids = self.extract_graphs(docx)

    def extract_graphs(self, docx):
        ids = []
        flag = False
        for i, p in enumerate(docx.paragraphs):
            if not flag and p.text.replace(" ", "").strip() == "目录":
                flag = True
            if flag:
                root = ET.fromstring(p._p.xml)
                pic_str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
                pics = root.findall(pic_str)
                image_str = "*/{urn:schemas-microsoft-com:vml}shape/{urn:schemas-microsoft-com:vml}imagedata"
                for pic in pics:
                    pict = pic.findall(image_str)
                    if len(pict) > 0:
                        text = docx.paragraphs[i + 1].text
                        if not text.endswith("示意图"):
                            ids.append(pict[0].attrib[
                                           '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'])
        return ids

    def parse_file(self, type_name, file_name):
        stage = None
        match = re.search("\d{3}", file_name)
        if match is not None:
            span = match.span()
            stage = file_name[span[0]: span[1]]
            stage = str(int(stage))

        GPRF_INTE = None
        if "-" in file_name:
            match = re.search("K\d\+\d{3}[-~](K\d\+)?\d{3}", file_name)
        if "～" in file_name:
            match = re.search("K\d\+\d{3}[～~](K\d\+)?\d{3}", file_name)
        if match is not None:
            span = match.span()
            GPRF_INTE = file_name[span[0]: span[1]]
            if "-" in GPRF_INTE:
                GPRF_INTE = GPRF_INTE.split("-")
                pre = GPRF_INTE[0][: 3]
                GPRF_INTE[1] = pre + GPRF_INTE[1]
                GPRF_INTE = "~".join(GPRF_INTE)
        # 文件命名不规范，不符合正则
        else:
            GPRF_INTE = ""

        prefix = util.map_prefix(util.parse_prefix(file_name))

        return type_name + prefix + stage + "期" + GPRF_INTE


class PicturePDF:
    def __init__(self, type_name, file_name, input_path):
        self.file = file_name
        self.directory = self.parse_file(type_name, file_name)
        self.pixes = self.extract_graphs(input_path)

    def extract_graphs(self, input_path):
        pixes = []
        # flag = False 不知道怎么判断图注
        # find图片
        checkXO = r"/Type(?= */XObject)"
        checkIM = r"/Subtype(?= */Image)"
        pdf = fitz.open(input_path)
        # 图片计数
        imgcount = 0
        # 获取对象数量长度
        lenXREF = pdf._getXrefLength()

        # 遍历每一个对象
        for i in range(1, lenXREF):
            # 定义对象字符串
            text = pdf._getXrefString(i)
            isXObject = re.search(checkXO, text)
            # 使用正则表达式查看是否是图片
            isImage = re.search(checkIM, text)
            # 如果不是对象也不是图片，则continue
            if not isXObject or not isImage:
                continue
            imgcount += 1
            # 根据索引生成图像对象
            pix = fitz.Pixmap(pdf, i)
            pixes.append(pix)

        return pixes

    def parse_file(self, type_name, file_name):
        stage = None
        match = re.search("\d{3}", file_name)
        if match is not None:
            span = match.span()
            stage = file_name[span[0]: span[1]]
            stage = str(int(stage))

        GPRF_INTE = None
        if "-" in file_name:
            match = re.search("K\d\+\d{3}[-~](K\d\+)?\d{3}", file_name)
        if "～" in file_name:
            match = re.search("K\d\+\d{3}[～~](K\d\+)?\d{3}", file_name)
        if match is not None:
            span = match.span()
            GPRF_INTE = file_name[span[0]: span[1]]
            if "-" in GPRF_INTE:
                GPRF_INTE = GPRF_INTE.split("-")
                pre = GPRF_INTE[0][: 3]
                GPRF_INTE[1] = pre + GPRF_INTE[1]
                GPRF_INTE = "~".join(GPRF_INTE)
        # 文件命名不规范，不符合正则
        else:
            GPRF_INTE = ""

        prefix = util.map_prefix(util.parse_prefix(file_name))

        return type_name + prefix + stage + "期" + GPRF_INTE


class Processor(FileProcessBasic):
    name = "TSP-S1S2标"

    def save(self, output, record):
        output_path = os.path.join(output, "TSP_S1S2.csv")
        header = record.dict.keys()
        util.check_output_file(output_path, header)

        with open(output_path, "a+", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, record.dict.keys())
            w.writerow(record.dict)

    def save_fig(self, base, pictures, docx):
        base = os.path.join(base, "图片数据")
        util.checkout_directory(base)
        pic_dir = os.path.join(base, pictures.directory)
        util.checkout_directory(pic_dir)
        processed_pics = set()
        for i, p_id in enumerate(pictures.picture_ids):
            if not processed_pics.__contains__(p_id):
                processed_pics.add(p_id)
            else:
                continue
            img = docx.part.related_parts[p_id]
            file_type = img.filename.split(".")[-1]
            with open(os.path.join(pic_dir, "{}.{}".format(str(i + 1), file_type)), "wb") as f:
                f.write(img.blob)

    def save_fig_PDF(self, base, pictures):
        base = os.path.join(base, "图片数据")
        util.checkout_directory(base)
        pic_dir = os.path.join(base, pictures.directory)
        util.checkout_directory(pic_dir)
        for i, pix in enumerate(pictures.pixes):
            new_name = "{}.png".format(i + 1)
            # 如果pix.n<5,可以直接存为PNG
            if pix.n < 5:
                path = os.path.join(pic_dir, new_name)
                pix.writePNG(path)
            # 否则先转换CMYK
            else:
                pix0 = fitz.Pixmap(fitz.csRGB, pix)
                pix0.writePNG(os.path.join(pic_dir, new_name))
                pix0 = None
            # 释放资源
            pix = None

    def run(self, input_path, output_path):
        files_to_process = set()
        files_to_transform = set()
        pdf_to_process = set()

        for file in os.listdir(input_path):
            absolute_file_path = os.path.join(input_path, file)
            if file.endswith(".doc"):
                files_to_transform.add(absolute_file_path)
            elif file.endswith(".docx"):
                files_to_process.add(absolute_file_path)
            elif file.endswith(".pdf"):
                pdf_to_process.add(file)

        files_to_delete = util.batch_doc_to_docx(files_to_transform)
        files_to_process = files_to_process.union(files_to_delete)

        # 处理Word
        for file in files_to_process:
            docx = Document(file)
            INTE = get_INTE(docx.tables[1])
            for i in range(len(INTE)):
                record = Record(docx, INTE[i])
                self.save(output_path, record)
            # 图片提取
            pics = Picture(Processor.name, file.split("\\")[-1], docx)
            self.save_fig(output_path, pics, docx)
            print("提取完成" + file)

        # 处理PDF
        for file in pdf_to_process:
            path = os.path.join(input_path, file)
            para_conclusion = ""
            para_suggestion = ""
            get_conclusion = False
            get_suggestion = False
            table_analysis = []
            table_result = []
            with plb.open(path) as pdf:
                for page in pdf.pages:
                    page_content = page.extract_text()
                    sentences = page_content.split("\n")

                    for s in sentences:
                        s = s.strip()
                        if s.startswith("第") and s.endswith("页"):
                            continue

                        if s.startswith("7.1"):
                            get_conclusion = True
                            continue
                        elif s.startswith("7.2"):
                            get_conclusion = False
                        if get_conclusion:
                            para_conclusion += s
                            continue

                        if s.startswith("7.2"):
                            get_suggestion = True
                            continue
                        elif s.startswith("8"):
                            get_suggestion = False
                            break
                        if get_suggestion:
                            para_suggestion += s
                        continue

                    if "力学指标汇总表" in page_content:
                        tables = page.extract_tables()
                        table_result = tables[0]
                    if "地质预报汇总表" in page_content:
                        tables = page.extract_tables()
                        table_analysis = tables[0]

            INTE_PDF = get_INTE_PDF(table_result)
            for i in range(len(INTE_PDF)):
                record = RecordPDF(pdf, INTE_PDF[i], table_result, table_analysis, para_conclusion)
                self.save(output_path, record)
            # PDF图片提取
            pics_PDF = PicturePDF(Processor.name, file.split("\\")[-1], path)
            self.save_fig_PDF(output_path, pics_PDF)
            print("提取完成PDF文件" + file)

        for file in files_to_delete:
            if os.path.exists(file):
                os.remove(file)


if __name__ == "__main__":
    test = Processor()
    inputpath = "E:/Education/409iS3/PDF/TSP"
    outputpath = "E:/Education/409iS3/PDF/output"
    test.run(inputpath, outputpath)
