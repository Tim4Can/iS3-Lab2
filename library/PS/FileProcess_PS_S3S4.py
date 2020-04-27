from docx import Document
import os
import csv
import re
from library.FileProcessBasic import FileProcessBasic
import xml.etree.cElementTree as ET
import util

class Record:
    def __init__(self, docx):
        name, GSI_INTE = self.get_cover(docx)
        GSI_CHAI, GSI_INTE = util.parse_GSI_CHAI_and_GSI_INTE(name, GSI_INTE)

        para_result, para_prediction,para_situation = self.locate_paragraph(docx)
        GSI_GPR = self.get_GSI_GPR(para_result)
        appendix = docx.tables[4]
        GSI_LITH = self.get_GSI_LITH(para_prediction)
        GSI_WEA = self.get_GSI_WEA(para_prediction)
        GSI_WATG = self.get_GSI_WATG(appendix)

        # 地下水状态描述
        GSI_WATE = self.get_GSI_WATE(para_prediction)

        # 岩层产状
        GSI_RKAT=self.get_GSI_RKAT(appendix)

        # 节理数
        GSI_JTNB=""

        # 节理倾角
        GSI_JTAG=self.get_GSI_JTAG(para_situation)

        # 完整性
        GSI_ITGT=self.get_GSI_ITGT(para_situation)

        # 完整性对应等级
        GSI_IGDG=self.get_GSI_IGDG(appendix)

        self.dict = {
            "掌子面桩号": GSI_CHAI,
            "桩号区间": GSI_INTE,
            "地质雷达描述": GSI_GPR,
            "地下水状态描述": GSI_WATE,
            "地下水对应等级": GSI_WATG,
            "岩性": GSI_LITH,
            "风化程度": GSI_WEA,
            "岩层产状": GSI_RKAT,
            "节理数":GSI_JTNB,
            "节理倾角":GSI_JTAG,
            "完整性":GSI_ITGT,
            "完整性对应等级":GSI_IGDG
        }

    def get_cover(self, docx):
        name, GSI_INTE = None, None
        for paragraph in docx.paragraphs:
            if paragraph.text.startswith("隧道名称："):
                name = paragraph.text.split("：")[1].strip()
            if paragraph.text.startswith("预报里程："):
                GSI_INTE = paragraph.text.split("：")[1].strip()
            if name is not None and GSI_INTE is not None:
                return name, GSI_INTE

    def locate_paragraph(self, docx):
        para_result = ""        # 6.2 探测结果
        para_prediction = ""    # 6.3 前方地质情况预测
        para_situation = ""    # 6.1 隧道掌子面地质情况
        for i, p in enumerate(docx.paragraphs):
            # print(i)
            # print(docx.paragraphs[i].text.strip())
            if p.text.startswith("6.1"):
                i += 1
                p = docx.paragraphs[i]
                while not p.text.startswith("6.2") or p.text == "\n":
                    if not p.text.startswith("图"):
                        para_situation += p.text
                    # print("situaiton:")
                    # print(para_situation)
                    i += 1
                    p = docx.paragraphs[i]

            elif p.text.startswith("6.2"):
                i += 1
                p = docx.paragraphs[i]
                while not p.text.startswith("6.3")  or p.text == "\n":
                    if not p.text.startswith("图"):
                        para_result += p.text
                    # print("result")
                    # print(para_result)
                    i += 1
                    p = docx.paragraphs[i]
                i -= 1
            elif p.text.startswith("6.3"):
                if not p.text.startswith("图"):
                    i += 1
                    p = docx.paragraphs[i]
                    while not p.text.startswith("7") :
                        para_prediction += p.text
                        i += 1
                        p = docx.paragraphs[i]
                    # print("prediction")
                    # print(para_prediction)
        # print("result"+para_result)
        # print("pre"+para_prediction)
        # print("sit"+para_situation)
        return para_result, para_prediction,para_situation

    # 掌子面桩号
    def get_GSI_CHAI(self, table):
        for i in range(len(table.rows)):
            tmp = list(table.rows[i].cells)
            cols = sorted(set(tmp), key=tmp.index)
            for j in range(len(cols)):
                if cols[j].text == '掌子面桩号' and j < len(cols) - 1:
                    GSI_CHAI = re.sub(u"\\（.*?）", "", cols[j + 1].text)
                    return GSI_CHAI

    # 桩号区间
    def get_GSI_INTE(self, para):
        GSI_INTE = ''
        for i in range(len(para)):
            if para[i] == '+':
                j = i
                while para[j] != '（':
                    j = j - 1
                while para[j + 1] != '）':
                    GSI_INTE = GSI_INTE + (para[j + 1])
                    j = j + 1
                break
        return GSI_INTE

    # 地质雷达描述
    def get_GSI_GPR(self, para):
        # 地质雷达描述
        GSI_GPR = "无"
        try:
            start = para.find("反射波的基本规律")
            start = para.find("：", start) + 1
            # end = para.find("反射频率", start)
            end = para.find("。", start)
            GSI_GPR = para[start: end]
            return GSI_GPR
        except:
            return GSI_GPR

    # 岩性
    def get_GSI_LITH(self, para):
        GSI_LITH=""
        start=para.find("岩性：")
        start=para.find("风化",start)
        end=para.find("。",start)
        GSI_LITH=para[start+len("风化"):end]
        if GSI_LITH == "":
            GSI_LITH = "无"
        return GSI_LITH

    # 风化程度
    def get_GSI_WEA(self, para):
        GSI_WEA = ""
        start = para.find("岩性：")
        end = para.find("风化", start)
        GSI_WEA = para[start + len("岩性："):end+len("风化")]
        if GSI_WEA == "":
            GSI_WEA = "无"
        return GSI_WEA

    # 层岩产状
    def get_GSI_RKAT(self,table):
        GSI_RKAT = ""
        for row in table.rows:
            if row.cells[0].text.strip() == "围岩岩性":
                content=row.cells[1].text.strip()
                content.replace("。","，")
                start=content.find("岩层产状")
                start=content.find("：",start)
                end=content.find("，",start)
                GSI_RKAT = content[start+len("："):end]
                # print(GSI_RKAT)
                if GSI_RKAT != "":
                    break

        if GSI_RKAT == "":
            GSI_RKAT = "无"
        return GSI_RKAT

    # 完整性
    def get_GSI_ITGT(self, para):
        start=para.find("总体")
        end=para.find("。",start)
        GSI_ITGT=para[start:end]
        if GSI_ITGT is None:
            GSI_ITGT="无"
        return GSI_ITGT

    # 节理倾角
    def get_GSI_JTAG(self, para):
        start=para.find("主要")
        end=para.find("。", start)
        GSI_JTAG=para[start:end]
        if GSI_JTAG is None:
            GSI_JTAG="无"
        return GSI_JTAG


    # 地下水状态描述
    def get_GSI_WATE(self,para):
        start=para.find("地下水：")
        end=para.find("。", start)
        GSI_WATE=para[start+len("地下水："):end]
        # print(GSI_WATE)
        if GSI_WATE is None:
            GSI_WATE="无"
        return GSI_WATE

    # 地下水对应等级
    def get_GSI_WATG(self, table):
        GSI_WATG = ""
        for row in table.rows:
            if row.cells[0].text.strip() == '地下水':
                watgs = set()
                for i in range(1, len(row.cells)):
                    cell = row.cells[i]
                    if "√" in cell.text:
                        watgs.add(cell.text.replace("√", "").strip())
                GSI_WATG = "~".join(watgs)
                # print(GSI_WATG)
                if GSI_WATG != "":
                    break

        if GSI_WATG == "":
            GSI_WATG = "无"
        return GSI_WATG

    # 完整性对应等级
    def get_GSI_IGDG(self,table):
        GSI_IGDG=""
        for row in table.rows:
            if row.cells[0].text.strip() == "完整性":
                content=row.cells[1].text.strip()
                content.replace("。","，")
                start=content.find("岩体")
                end=content.find("，",start)
                GSI_IGDG = content[start:end]
                # print(GSI_IGDG)
                if GSI_IGDG != "":
                    break

        if GSI_IGDG == "":
            GSI_IGDG = "无"
        return GSI_IGDG

class Picture:
    def __init__(self, type_name, file_name, docx):
        self.file = file_name
        self.directory = self.parse_file(type_name, file_name)
        self.picture_ids = self.extract_graphs(docx)

    def extract_graphs(self, docx):
        ids = []
        flag = False
        for i, p in enumerate(docx.paragraphs):
            if not flag and p.text.replace(" ", "").strip() == "目录":
                flag = True
            if flag:
                root = ET.fromstring(p._p.xml)
                pic_str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
                pics = root.findall(pic_str)
                image_str = "*/{urn:schemas-microsoft-com:vml}shape/{urn:schemas-microsoft-com:vml}imagedata"
                for pic in pics:
                    pict = pic.findall(image_str)
                    if len(pict) > 0:
                        text = docx.paragraphs[i + 1].text
                        if not text.endswith("示意图"):
                
                          ids.append(pict[0].attrib[
                                           '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'])
        return ids

    def parse_file(self, type_name, file_name):
        stage = None
        match = re.search("\d{3}", file_name)

        if match is not None:
            span = match.span()
            stage = file_name[span[0]: span[1]]
            stage = str(int(stage))

        GSI_INTE = None
        match = re.search("K\d\+\d{3}[～~](K\d\+)?\d{3}", file_name)

        if match is not None:
            span = match.span()
            GSI_INTE = file_name[span[0]: span[1]]
            if "-" in GSI_INTE:
                GSI_INTE = GSI_INTE.split("-")
                pre = GSI_INTE[0][: 3]
                GSI_INTE[1] = pre + GSI_INTE[1]
                GSI_INTE = "~".join(GSI_INTE)

        prefix = util.map_prefix(util.parse_prefix(file_name))

        return type_name + prefix + stage + "期" + GSI_INTE




class Processor(FileProcessBasic):
    name = "S3S4标"

    def save(self, output, record):
        output_path = os.path.join(output, "PS_S3S4.csv")
        header = record.dict.keys()
        util.check_output_file(output_path, header)

        with open(output_path, "a+", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, record.dict.keys())
            w.writerow(record.dict)

    def save_fig(self, base, pictures, docx):
        base = os.path.join(base, "图片数据")
        util.checkout_directory(base)
        pic_dir = os.path.join(base, pictures.directory)
        util.checkout_directory(pic_dir)
        processed_pics = set()
        for i, p_id in enumerate(pictures.picture_ids):
            if not processed_pics.__contains__(p_id):
                processed_pics.add(p_id)
            else:
                continue
            img = docx.part.related_parts[p_id]
            file_type = img.filename.split(".")[-1]
            with open(os.path.join(pic_dir, "{}.{}".format(str(i + 1), file_type)), "wb") as f:
                f.write(img.blob)


    def run(self, input_path, output_path):
        files_to_process = set()
        files_to_transform = set()
        for file in os.listdir(input_path):
            absolute_file_path = os.path.join(input_path, file)
            if file.endswith(".doc"):
                files_to_transform.add(absolute_file_path)
            elif file.endswith(".docx"):
                files_to_process.add(absolute_file_path)
        files_to_delete = util.batch_doc_to_docx(files_to_transform)
        files_to_process = files_to_process.union(files_to_delete)

        for file in files_to_process:
            docx = Document(file)
            record = Record(docx)
            self.save(output_path, record)

            pics = Picture(Processor.name, file.split("\\")[-1], docx)
            self.save_fig(output_path, pics, docx)           
            print("提取完成" + file)

        for file in files_to_delete:
            if os.path.exists(file):
                os.remove(file)


if __name__ == "__main__":
    test = Processor()

    inputpath = "C:/Users/DELL/Desktop/iS3/新建文件夹"
    outputpath = "C:/Users/DELL/Desktop/iS3"
    test.run(inputpath, outputpath)
