from docx import Document
import os
import re
import xml.etree.cElementTree as ET
from library.FileProcessBasic import FileProcessBasic
import util
class Picture:
    def __init__(self, inte, id, type):
        self.inte = inte
        self.id = id
        self.type = type

class PictureProcess:
    def __init__(self, type_name, file_name, docx):
        self.file = file_name
        self.directory = self.parse_file(type_name, file_name)
        self.picture_ids = self.extract_graphs(docx)

    def extract_graphs(self, docx):
        pictures = []
        flag = False
        for i, p in enumerate(docx.paragraphs):
            if not flag and p.text.replace(" ", "").strip() == "目录":
                flag = True
            if flag:
                root = ET.fromstring(p._p.xml)
                pic_str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
                pics = root.findall(pic_str)
                image_str = "*/{urn:schemas-microsoft-com:vml}shape/{urn:schemas-microsoft-com:vml}imagedata"
                for pic in pics:
                    pict = pic.findall(image_str)
                    if len(pict) > 0:
                        text = docx.paragraphs[i + 1].text
                        start = text.find("ZK")
                        if start < 0:
                            start = text.find("YK")
                        if start < 0:
                            start = text.find("K")
                        if text.endswith("掌子面地质素描图"):

                            end = text.find("掌子面地质素描图")
                            picture = Picture(text[start:end],pict[0].attrib[
                                           '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'],
                                              "TSSI")

                            pictures.append(picture)
                        elif text.endswith("掌子面照片"):
                            end = text.find("掌子面照片")
                            picture = Picture(text[start:end], pict[0].attrib[
                                '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'],
                                              "TSIM")
                            pictures.append(picture)
        return pictures

    def parse_file(self, type_name, file_name):
        stage = None
        match = re.search("\d{3}", file_name)
        if match is not None:
            span = match.span()
            stage = file_name[span[0]: span[1]]
            stage = str(int(stage))

        GSI_INTE = None
        match = re.search("K\d\+\d{3}[-～](K\d\+)?\d{3}", file_name)
        if match is not None:
            span = match.span()
            GSI_INTE = file_name[span[0]: span[1]]
            if "-" in GSI_INTE:
                GSI_INTE = GSI_INTE.split("-")
                pre = GSI_INTE[0][: 3]
                GSI_INTE[1] = pre + GSI_INTE[1]
                GSI_INTE = "~".join(GSI_INTE)
        else:
            GSI_INTE=""

        prefix = util.map_prefix(util.parse_prefix(file_name))

        return type_name + prefix + stage + "期" + GSI_INTE

class Processor(FileProcessBasic):
    name = "PS-S1S2标"

    def save(self, output, record):
        print("")


    def save_fig(self, base, pictures, docx):
        base = os.path.join(base, "pic")
        util.checkout_directory(base)
        pic_dir = os.path.join(base, pictures.directory)
        util.checkout_directory(pic_dir)
        processed_pics = set()
        i = 0
        j = 0
        for p in pictures.picture_ids:
            p_id = p.id
            if not processed_pics.__contains__(p_id):
                processed_pics.add(p_id)
            else:
                continue
            img = docx.part.related_parts[p_id]
            file_type = img.filename.split(".")[-1]
            if p.type == "TSSI":
                pic_name=p.inte + "_" + p.type + "_" + str(i+1)
                with open(os.path.join(pic_dir, "{}.{}".format(pic_name, file_type)), "wb") as f:
                    f.write(img.blob)
                i=+1
            elif p.type == "TSIM":
                pic_name = p.inte + "_" + p.type + "_" + str(j + 1)
                with open(os.path.join(pic_dir, "{}.{}".format(pic_name, file_type)), "wb") as f:
                    f.write(img.blob)
                j=+1



    def run(self, input_path, output_path):
        files_to_process = set()
        files_to_transform = set()

        for file in os.listdir(input_path):
            absolute_file_path = os.path.join(input_path, file)
            if file.endswith(".doc"):
                files_to_transform.add(absolute_file_path)
            elif file.endswith(".docx"):
                files_to_process.add(absolute_file_path)
        files_to_delete = util.batch_doc_to_docx(files_to_transform)
        files_to_process = files_to_process.union(files_to_delete)

        for file in files_to_process:
            docx = Document(file)


            pics = PictureProcess(Processor.name, file.split("\\")[-1], docx)
            self.save_fig(output_path, pics, docx)
            print("提取完成" + file)


        for file in files_to_delete:
            if os.path.exists(file):
                os.remove(file)
if __name__ == "__main__":
    test = Processor()

    inputpath = "/Users/budi/Desktop/iS3/PS"
    outputpath = "/Users/budi/Desktop/iS3/OUTPUT"
    test.run(inputpath, outputpath)