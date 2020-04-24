from docx import Document
import os
import csv
import re
import xml.etree.cElementTree as ET
from library.FileProcessBasic import FileProcessBasic
import util


class Record:
    def __init__(self, docx):
        name, SKTH_INTE = self.get_cover(docx)
        SKTH_CHAI, SKTH_INTE = util.parse_SKTH_CHAI_and_SKTH_INTE(name, SKTH_INTE)

        # 文字中提取
        para_result, para_conclusion, para_suggestion = self.locate_paragraph(docx)

        SKTH_STRU = self.get_SKTH_STRU(para_suggestion)
        SKTH_STAB = self.get_SKTH_STAB(para_conclusion)
        SKTH_SURR = self.get_SKTH_SURR(para_conclusion)
        SKTH_INTE2 = self.get_SKTH_INTE2(para_conclusion)

        # 附件表格中提取
        appendix = docx.tables[2]

        SKTH_WEA = self.get_SKTH_WEA(appendix)
        SKTH_LITH = self.get_SKTH_LITH(appendix)
        SKTH_FAUL = self.get_SKTH_FAUL(appendix)
        SKTH_WATG = self.get_SKTH_WATG(appendix)
        SKTH_FORM = self.get_SKTH_FORM(appendix)
        SKTH_STRE = self.get_SKTH_STRE(appendix)

        # 未找到对应描述
        SKTH_WATE = self.get_SKTH_WATE()
        SKTH_JOIQ = self.get_SKTH_JOIQ()
        SKTH_JOIN = self.get_SKTH_JOIN()
        SKTH_INTG = self.get_SKTH_INTG()

        self.dict = {
            "掌子面桩号": SKTH_CHAI,
            "桩号区间": SKTH_INTE,
            "地下水状态描述": SKTH_WATE,
            "地下水对应等级": SKTH_WATG,
            "岩性": SKTH_LITH,
            "岩层产状": SKTH_FORM,
            "风化程度": SKTH_WEA,
            "节理数": SKTH_JOIQ,
            "节理倾角": SKTH_JOIN,
            "完整性": SKTH_INTE2,
            "完整性对应等级": SKTH_INTG,
            "围岩级别": SKTH_SURR,
            "结构面形状": SKTH_STRU,
            "断层": SKTH_FAUL,
            "高应力、特殊地质": SKTH_STRE,
            "围岩稳定情况": SKTH_STAB,
        }

    def get_cover(self, docx):
        name, SKTH_INTE = None, None
        for paragraph in docx.paragraphs:
            if paragraph.text.startswith("隧道名称："):
                name = paragraph.text.split("：")[1].strip()
            if paragraph.text.startswith("预报里程："):
                SKTH_INTE = paragraph.text.split("：")[1].strip()
            if name is not None and SKTH_INTE is not None:
                return name, SKTH_INTE

    def locate_paragraph(self, docx):

        para_result = ""  # 6.2 探测结果
        para_conclusion = ""  # 7.1 结论
        para_suggestion = ""  # 7.2 建议
        for i, p in enumerate(docx.paragraphs):
            if p.text.startswith("6.2"):
                i += 1
                p = docx.paragraphs[i]
                while not p.text.startswith("7"):
                    if not p.text.startswith("图"):
                        para_result += p.text
                    i += 1
                    p = docx.paragraphs[i]
            elif p.text.startswith("7.1"):
                i += 1
                p = docx.paragraphs[i]
                while not p.text.startswith("7.2"):
                    para_conclusion += p.text
                    i += 1
                    p = docx.paragraphs[i]
                i -= 1
            elif p.text.startswith("7.2"):
                i += 1
                p = docx.paragraphs[i]
                while not p.text.startswith("附件") or p.text == "\n":
                    para_suggestion += p.text
                    i += 1
                    p = docx.paragraphs[i]
        return para_result, para_conclusion, para_suggestion

    # 掌子面桩号
    def get_SKTH_CHAI(self, table):
        for i in range(len(table.rows)):
            tmp = list(table.rows[i].cells)
            cols = sorted(set(tmp), key=tmp.index)
            for j in range(len(cols)):
                if cols[j].text == '掌子面桩号' and j < len(cols) - 1:
                    SKTH_CHAI = re.sub(u"\\（.*?）", "", cols[j + 1].text)
                    return SKTH_CHAI

    # 桩号区间
    def get_SKTH_INTE(self, para):
        SKTH_INTE = ''
        for i in range(len(para)):
            if para[i] == '+':
                j = i
                while para[j] != '（':
                    j = j - 1
                while para[j + 1] != '）':
                    SKTH_INTE = SKTH_INTE + (para[j + 1])
                    j = j + 1
                break
        return SKTH_INTE

    # 岩性
    def get_SKTH_LITH(self, table):
        SKTH_LITH = ""
        for row in table.rows:
            if row.cells[0].text.strip().replace(" ", "") == "岩性":
                SKTH_LITH = row.cells[4].text
        if SKTH_LITH == "":
            SKTH_LITH = "无"
        return SKTH_LITH

    # 岩层产状
    def get_SKTH_FORM(self, table):
        SKTH_FORM = ""
        for row in table.rows:
            for i in range(1, len(row.cells)):
                if row.cells[i].text.strip() == "岩层产状":
                    while True:
                        i += 1
                        if row.cells[i].text.strip() != "岩层产状":
                            break
                    SKTH_FORM = row.cells[i].text.strip()
                    break
        if SKTH_FORM == "\\" or "":
            SKTH_FORM = "无"
        return SKTH_FORM

    # 风化程度
    def get_SKTH_WEA(self, table):
        SKTH_WEA = ""
        for row in table.rows:
            if row.cells[0].text.strip().replace(" ", "") == "风化程度":
                weas = set()
                for i in range(1, len(row.cells)):
                    cell = row.cells[i]
                    if "√" in cell.text:
                        weas.add(cell.text.replace("√", "").strip())
                SKTH_WEA = "~".join(weas) + "风化"
                break
        if SKTH_WEA == "":
            SKTH_WEA = "无"
        return SKTH_WEA

    # 节理数
    def get_SKTH_JOIQ(self):
        return "无"

    # 节理倾角
    def get_SKTH_JOIN(self):
        return "无"

    # 完整性
    def get_SKTH_INTE2(self, para):
        start = para.find("围岩整体")
        end = para.find("。", start)
        SKTH_INTE2 = para[start: end].replace("及稳定性", "")
        if SKTH_INTE2 is None:
            SKTH_INTE2 = "无"
        return SKTH_INTE2

    # 完整性对应等级
    def get_SKTH_INTG(self):
        return "无"

    # 结构面形状
    def get_SKTH_STRU(self, para):
        start = para.find("呈")
        end = para.find("结构", start) + 2
        SKTH_STRU = para[start: end]
        if SKTH_STRU is None:
            SKTH_STRU = "无"
        return SKTH_STRU

    # 稳定性
    def get_SKTH_STAB(self, para):
        start = para.find("稳定性")
        end = para.find("。", start)
        SKTH_STAB = para[start: end]
        if SKTH_STAB is None:
            SKTH_STAB = "无"
        return SKTH_STAB

    # 围岩级别   ！暂且按照设计围岩等级取
    def get_SKTH_SURR(self, para):
        keywords = "设计围岩等级为"
        start = para.find(keywords) + len(keywords)
        end = para.find("级", start)
        SKTH_SURR = para[start: end]
        if SKTH_SURR is None:
            SKTH_SURR = "无"
        return SKTH_SURR

    # 地下水状态描述
    def get_SKTH_WATE(self):
        return "无"

    # 地下水对应等级
    def get_SKTH_WATG(self, table):
        SKTH_WATG = ""
        for row in table.rows:
            if row.cells[0].text.strip() == '地下水状态':
                watgs = set()
                for i in range(1, len(row.cells)):
                    cell = row.cells[i]
                    if "√" in cell.text:
                        watgs.add(cell.text.replace("√", "").strip())
                SKTH_WATG = "~".join(watgs)
                break
        if SKTH_WATG == "":
            SKTH_WATG = "无"
        return SKTH_WATG

    # 断层   结构需要优化
    def get_SKTH_FAUL(self, table):
        SKTH_FAUL = ""
        result = []
        for row in table.rows:
            for i in range(1, len(row.cells)):
                if row.cells[0].text.strip() == "断层":
                    while True:
                        i += 1
                        if row.cells[i].text.strip() != "断层":
                            break
                    while True:
                        i += 1
                        if row.cells[i].text.strip() != "断层产状":
                            break
                    result.append(row.cells[i].text.strip())
                    while True:
                        i += 1
                        if row.cells[i].text.strip() != "断层宽度（m）":
                            break
                    result.append(row.cells[i].text.strip())
                    while True:
                        i += 1
                        if row.cells[i].text.strip() != "断层性质":
                            break
                    result.append(row.cells[i].text.strip())
                    SKTH_FAUL = "".join(result)
                    break
        if SKTH_FAUL == "":
            SKTH_FAUL = "无"
        return SKTH_FAUL

    # 高应力、特殊地质
    def get_SKTH_STRE(self, table):
        SKTH_STRE = ""
        for row in table.rows:
            if row.cells[0].text.strip() == '初始应力状态':
                stres = set()
                for i in range(1, len(row.cells)):
                    cell = row.cells[i]
                    if "√" in cell.text:
                        stres.add(cell.text.replace("√", "").strip())
                SKTH_STRE = "~".join(stres)
                break
        if SKTH_STRE == "":
            SKTH_STRE = "无"
        return SKTH_STRE


class Picture:
    def __init__(self, type_name, file_name, docx):
        self.file = file_name
        self.directory = self.parse_file(type_name, file_name)
        self.picture_ids = self.extract_graphs(docx)

    def extract_graphs(self, docx):
        ids = []
        flag = False
        for i, p in enumerate(docx.paragraphs):
            if not flag and p.text.replace(" ", "").strip() == "目录":
                flag = True
            if flag:
                root = ET.fromstring(p._p.xml)
                pic_str = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
                pics = root.findall(pic_str)
                image_str = "*/{urn:schemas-microsoft-com:vml}shape/{urn:schemas-microsoft-com:vml}imagedata"
                for pic in pics:
                    pict = pic.findall(image_str)
                    if len(pict) > 0:
                        text = docx.paragraphs[i + 1].text
                        if not text.endswith("示意图"):
                            ids.append(pict[0].attrib[
                                           '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'])
        return ids

    def parse_file(self, type_name, file_name):
        stage = None
        match = re.search("\d{3}", file_name)
        if match is not None:
            span = match.span()
            stage = file_name[span[0]: span[1]]
            stage = str(int(stage))

        SKTH_INTE = None
        match = re.search("K\d\+\d{3}[-~](K\d\+)?\d{3}", file_name)
        if match is not None:
            span = match.span()
            SKTH_INTE = file_name[span[0]: span[1]]
            if "-" in SKTH_INTE:
                SKTH_INTE = SKTH_INTE.split("-")
                pre = SKTH_INTE[0][: 3]
                SKTH_INTE[1] = pre + SKTH_INTE[1]
                SKTH_INTE = "~".join(SKTH_INTE)

        prefix = util.map_prefix(util.parse_prefix(file_name))

        return type_name + prefix + stage + "期" + SKTH_INTE


class Processor(FileProcessBasic):
    name = "S1S2标"

    def save(self, output, record):
        output_path = os.path.join(output, "PS_S1S2.csv")
        header = record.dict.keys()
        util.check_output_file(output_path, header)

        with open(output_path, "a+", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, record.dict.keys())
            w.writerow(record.dict)

    def save_fig(self, base, pictures, docx):
        base = os.path.join(base, "图片数据")
        util.checkout_directory(base)
        pic_dir = os.path.join(base, pictures.directory)
        util.checkout_directory(pic_dir)
        processed_pics = set()
        for i, p_id in enumerate(pictures.picture_ids):
            if not processed_pics.__contains__(p_id):
                processed_pics.add(p_id)
            else:
                continue
            img = docx.part.related_parts[p_id]
            file_type = img.filename.split(".")[-1]
            with open(os.path.join(pic_dir, "{}.{}".format(str(i + 1), file_type)), "wb") as f:
                f.write(img.blob)

    def run(self, input_path, output_path):
        files_to_process = set()
        files_to_transform = set()
        for file in os.listdir(input_path):
            absolute_file_path = os.path.join(input_path, file)
            if file.endswith(".doc"):
                files_to_transform.add(absolute_file_path)
            elif file.endswith(".docx"):
                files_to_process.add(absolute_file_path)
        files_to_delete = util.batch_doc_to_docx(files_to_transform)
        files_to_process = files_to_process.union(files_to_delete)

        for file in files_to_process:
            docx = Document(file)
            record = Record(docx)
            self.save(output_path, record)
            # 图片提取
            pics = Picture(Processor.name, file.split("\\")[-1], docx)
            self.save_fig(output_path, pics, docx)

            print("提取完成" + file)

        for file in files_to_delete:
            if os.path.exists(file):
                os.remove(file)


if __name__ == "__main__":
    test = Processor()
    inputpath = "E:/Education/409iS3/task/task2"
    outputpath = "E:/Education/409iS3/task/task2"
    test.run(inputpath, outputpath)
